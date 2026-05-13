"""
Targeted edits to Chapter 3 of Pulsefy_Diploma_Thesis.docx.

Applies four corrections grounded in what actually exists (or is planned)
in the codebase:

  1. 3.1.1 — rewrite the Administrator paragraph: the only documented
     admin privilege is bypassing the subscription tier check so the
     account can test premium AI models without paying.

  2. 3.1.3 — remove the invented quota lines from UC4, UC5, UC7
     (no daily/monthly quotas exist in the code).

  3. 3.1.3.11 — rename UC11 to "Access AI models without subscription"
     and replace its lettered items with three claims that match the
     real (planned) admin functionality.

  4. 3.1.4.11 — same rename. Update the description-table cells, the
     table caption, and the figure caption.

The use case diagram (Figure 8) and UC11's SSD (Figure 19) need to be
re-drawn in Mermaid by the user — those are images, not text.
"""

from docx import Document
from docx.shared import Pt

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def replace_paragraph_text(p, new_text, fontsize=None):
    """Replace paragraph text by clearing existing runs and adding one new run.

    Paragraph-level style (e.g. Heading 3) is preserved on the new run.
    """
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    new_run = p.add_run(new_text)
    if fontsize is not None:
        new_run.font.size = Pt(fontsize)


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def set_cell_text_plain(cell, text, bold=False):
    """Replace a cell's content with a single paragraph of `text`."""
    # Remove all existing paragraphs except the first
    paragraphs = list(cell.paragraphs)
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    # Clear runs from the first paragraph
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    run = first.add_run(text)
    if bold:
        run.bold = True


def set_cell_multiline(cell, lines, bold=False):
    """Replace a cell's content with multiple paragraphs (one per line)."""
    paragraphs = list(cell.paragraphs)
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    # First line into the first paragraph
    if lines:
        run = first.add_run(lines[0])
        if bold:
            run.bold = True
        # Add remaining lines as new paragraphs
        for line in lines[1:]:
            new_p = cell.add_paragraph()
            run = new_p.add_run(line)
            if bold:
                run.bold = True


doc = Document(DOCX)

# ---------------------------------------------------------------
# 1. 3.1.1  --  rewrite the Administrator paragraph
# ---------------------------------------------------------------
new_admin_para = (
    "The system also has one agent: the administrator. An administrator "
    "is an authenticated user whose account is flagged with an "
    "administrator role in the database, rather than through any in-app "
    "registration flow. The administrator's only privilege is access to "
    "all AI generation models for testing purposes: requests issued by an "
    "administrator bypass the subscription tier check, which lets the "
    "administrator generate using the premium MusicGen checkpoint and "
    "other paid models without an active premium subscription. The "
    "administrator has no user-management, content-moderation, or "
    "system-monitoring functionality; those are out of scope for the "
    "current release."
)

found_admin = False
for p in doc.paragraphs:
    if p.text.startswith("The system also has one agent"):
        replace_paragraph_text(p, new_admin_para)
        found_admin = True
        break
print(f"3.1.1 admin paragraph updated: {found_admin}")

# ---------------------------------------------------------------
# 2. 3.1.3  --  remove invented quota / timing lines
# ---------------------------------------------------------------
QUOTA_SIGNATURES = [
    "free-tier limit on the number of images that can be generated per day",
    "per-month generation limit on the free tier",
    "estimated generation time before starting",
    "free-tier limit on the number of voiceovers generated per day",
]
removed = []
for p in list(doc.paragraphs):
    if any(sig in p.text for sig in QUOTA_SIGNATURES):
        removed.append(p.text[:80])
        remove_paragraph(p)
print(f"3.1.3 quota lines removed: {len(removed)}")
for r in removed:
    print(f"  - {r}")

# ---------------------------------------------------------------
# 3. 3.1.3.11  --  rewrite UC11 heading and lettered items
# ---------------------------------------------------------------
new_uc11_heading = "3.1.3.11 Access AI models without subscription – UC11"
new_uc11_items = [
    "a) System recognises accounts with the administrator flag in the user record – 10",
    "b) System exempts administrator accounts from the subscription tier check on AI generation requests – 10",
    "c) System grants administrators access to the premium MusicGen checkpoint regardless of subscription tier – 10",
]

# Find old UC11 heading
old_heading = None
for p in doc.paragraphs:
    if "3.1.3.11 Manage users and content" in p.text:
        old_heading = p
        break

if old_heading is None:
    print("WARN: 3.1.3.11 heading not found")
else:
    # Update heading text (style preserved)
    replace_paragraph_text(old_heading, new_uc11_heading)

    # Walk forward through paragraphs after the heading until next Heading 3.
    # Compare by XML element identity since python-docx rebuilds the wrapper list.
    target_elem = old_heading._element
    paras = list(doc.paragraphs)
    heading_idx = next(
        i for i, p in enumerate(paras) if p._element is target_elem
    )
    item_paras = []
    for q in paras[heading_idx + 1:]:
        if q.style.name.startswith("Heading"):
            break
        if q.text.strip() and q.text.strip()[:2] in [f"{c})" for c in "abcdefghij"]:
            item_paras.append(q)
        elif q.text.strip().startswith(tuple(f"{c})" for c in "abcdefghij")):
            item_paras.append(q)

    # Replace existing item texts; remove extras
    for i, item_para in enumerate(item_paras):
        if i < len(new_uc11_items):
            replace_paragraph_text(item_para, new_uc11_items[i])
        else:
            remove_paragraph(item_para)
    print(f"3.1.3.11 items updated: {len(item_paras)} found, "
          f"{min(len(item_paras), len(new_uc11_items))} kept and rewritten, "
          f"{max(0, len(item_paras) - len(new_uc11_items))} removed")

# ---------------------------------------------------------------
# 4. 3.1.4.11  --  rename heading, update table, update captions
# ---------------------------------------------------------------
NEW_NAME = "Access AI models without subscription"

# 4a. Heading
for p in doc.paragraphs:
    if "3.1.4.11" in p.text and "Manage users and content" in p.text:
        replace_paragraph_text(p, f"3.1.4.11 UC11 {NEW_NAME}")
        break

# 4b. Description table
target_table = None
for t in doc.tables:
    second_col = [r.cells[1].text for r in t.rows]
    if any("Manage users and content" in c for c in second_col):
        target_table = t
        break

if target_table is None:
    print("WARN: UC11 description table not found")
else:
    # Six rows: Use Case Name / Scope / Actor / Preconditions / Post-conditions / Main success scenario
    new_rows = [
        ("Use Case Name:", f"UC11 {NEW_NAME}", False),
        ("Scope:", "System", False),
        ("Actor:", "Administrator", False),
        ("Preconditions:", "The user is authenticated and the user record has the administrator flag set.", False),
        ("Post-conditions:", "The user can generate ad assets through any AI model (premium or free) without a subscription tier check.", False),
        ("Main success scenario:", None, True),  # multiline filled below
    ]
    main_scenario_lines = [
        "1. The administrator opens any AI generation panel.",
        "2. The system shows the form for the chosen subsystem.",
        "3. The administrator submits the generation request.",
        "4. The system checks the administrator flag on the account.",
        "5. The system bypasses the subscription tier check and invokes the requested model.",
        "6. The system stores the generated asset and displays it.",
    ]

    for i, (label, value, is_multiline) in enumerate(new_rows):
        if i >= len(target_table.rows):
            break
        row = target_table.rows[i]
        set_cell_text_plain(row.cells[0], label, bold=True)
        if is_multiline:
            set_cell_multiline(row.cells[1], main_scenario_lines, bold=False)
        else:
            set_cell_text_plain(row.cells[1], value, bold=False)
    print("3.1.4.11 description table updated")

# 4c. Captions
for p in doc.paragraphs:
    if p.text.startswith("Table 12: UC11 Manage users and content"):
        replace_paragraph_text(p, f"Table 12: UC11 {NEW_NAME} Description", fontsize=10)
    elif p.text.startswith("Figure 19: UC11 Manage users and content"):
        replace_paragraph_text(p, f"Figure 19: UC11 {NEW_NAME} system sequence diagram", fontsize=10)
    elif "[FIGURE 19" in p.text and "Manage users and content" in p.text:
        replace_paragraph_text(p, f"[FIGURE 19: UC11 {NEW_NAME} system sequence diagram — to be drawn]", fontsize=10)

doc.save(DOCX)
print(f"Saved: {DOCX}")
