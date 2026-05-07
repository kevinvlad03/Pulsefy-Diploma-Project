"""
Targeted edit: replace the four red footnote markers in Chapter 1 with
real source info, and soften the F3 prose claim that has no authoritative
single source for the €1k–€10k range.

Verified 2026-05-07 via WebFetch + WebSearch.
"""

from docx import Document
from docx.shared import Pt, RGBColor

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def find_para(doc, signature):
    for i, p in enumerate(doc.paragraphs):
        if signature in p.text:
            return i
    return -1


def rewrite_paragraph(p, prose, marker_text):
    """Replace paragraph contents with prose + red-italic marker."""
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    main = p.add_run(prose)
    main.font.size = Pt(11)
    if marker_text:
        m = p.add_run(f"  [{marker_text}]")
        m.italic = True
        m.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        m.font.size = Pt(10)


doc = Document(DOCX)

# ----- 1.1 P1 (F1) -----
i = find_para(doc, "TikTok, Instagram Reels, and YouTube Shorts each report")
assert i >= 0, "could not find 1.1 P1 (TikTok / Reels / Shorts paragraph)"
rewrite_paragraph(
    doc.paragraphs[i],
    "Short-form vertical video has become the dominant format for "
    "social-media advertising. TikTok, Instagram Reels, and YouTube Shorts "
    "each report user audiences in the hundreds of millions. eMarketer "
    "estimated TikTok's monthly active users at approximately 955 million "
    "for 2025, building from the one billion mark ByteDance officially "
    "announced in September 2021. Brands and creators that need to reach "
    "those audiences are concentrating their advertising effort on this "
    "format because it matches where viewer attention has moved.",
    "footnote F1: TikTok MAU per eMarketer (via Statista, 2025); ByteDance "
    "official 1 B announcement (Sept 2021). URL: "
    "https://www.statista.com/statistics/1327116/number-of-global-tiktok-users/ "
    "— add Word footnote here",
)

# ----- 1.1 P2 (F2) -----
i = find_para(doc, "advertising spend has followed the attention")
assert i >= 0, "could not find 1.1 P2 (digital ad spend paragraph)"
rewrite_paragraph(
    doc.paragraphs[i],
    "The advertising spend has followed the attention. WARC's September 2025 "
    "global forecast estimated worldwide ad spending at 1.17 trillion US "
    "dollars in 2025, a 7.4 percent year-on-year increase, with social "
    "media accounting for 26.2 percent of the total at 306.4 billion "
    "dollars and growing 14.9 percent year on year. Social platforms now "
    "capture the majority of incremental ad dollars, and short-form video "
    "is among the fastest-growing single formats inside that share.",
    "footnote F2: WARC Global Ad Spend forecast Sept 2025 via eMarketer. "
    "URL: https://www.emarketer.com/content/global-ad-spending-now-"
    "estimated-rise-7-4--warc-s-first-positive-revision-over-year — add "
    "Word footnote here",
)

# ----- 1.3 P1 (F3) -- soften the claim because no authoritative single source
#       supports the exact €1 k–€10 k range. -----
i = find_para(doc, "absolute cost of creative-agency work")
assert i >= 0, "could not find 1.3 P1 (agency cost paragraph)"
rewrite_paragraph(
    doc.paragraphs[i],
    "The cost gap between what small advertisers need and what is available "
    "to them has two layers. The first is the absolute cost of "
    "creative-agency work, which typically runs into the thousands of "
    "euros per campaign for the kind of small project a side-hustle "
    "creator might commission. Romanian and broader Eastern European "
    "agencies generally bill at lower hourly rates than Western European "
    "or North American studios, but a single short-form ad campaign still "
    "represents a recurring expense that a retailer with a few hundred "
    "euros in monthly ad budget cannot absorb.",
    "footnote F3: Romanian agency hourly rates and campaign-cost range — "
    "verify via Sortlist Romania directory "
    "(https://www.sortlist.com/advertising/romania-ro) and Clutch.co "
    "agency directory; the specific euro range was softened on 2026-05-07 "
    "because no single authoritative source publishes it",
)

# ----- 1.3 P2 (F4) -----
i = find_para(doc, "Adobe Creative Cloud, the industry-standard package")
assert i >= 0, "could not find 1.3 P2 (Adobe pricing paragraph)"
rewrite_paragraph(
    doc.paragraphs[i],
    "The second layer is the subscription cost of professional production "
    "software. Adobe Creative Cloud, the industry-standard package for "
    "design, image editing, and video, is offered as the Creative Cloud "
    "Pro plan (the successor to the previous All Apps tier introduced in "
    "August 2025) at 69.99 US dollars per month for individual users, "
    "with regional pricing in euro and Romanian leu published on Adobe's "
    "national pages. Final Cut Pro, DaVinci Resolve Studio, Logic Pro, "
    "and similar professional tools each carry one-time or subscription "
    "fees, on top of which most have steep learning curves measured in "
    "months rather than days. The total stack is unattractive for "
    "someone producing a handful of ads per month.",
    "footnote F4: Adobe Creative Cloud Pro pricing — adobe.com/creativecloud/pro.html "
    "(USD); fetch adobe.com/ro/creativecloud/plans.html for the current "
    "RON price before final submission",
)

doc.save(DOCX)
print(f"Saved: {DOCX}")
