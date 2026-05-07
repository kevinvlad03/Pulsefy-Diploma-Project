"""
Pulsefy Diploma Thesis — skeleton builder (v2, ad-creation framing).

Reframed 2026-05-07: Pulsefy is an AI-assisted ad creation platform for
solo creators and small businesses producing short-form social-video ads.
Music streaming, MusicGen, gTTS, Pollinations imagegen, and the video
overlay pipeline are components of an ad-creation flow, not standalone
products.

Produces docs/Pulsefy_Diploma_Thesis.docx with:
- Three title pages (English bare, English with logo placeholders, Romanian
  with logo placeholders) using the ad-creation title
- Diploma Project Theme form (8 fields, ad-creation positioning,
  Iuliana Marin coordinator)
- Academic Honesty Statement
- TOC + Table of Figures placeholders
- Chapter 1 fully drafted across 5 subsections (humanizer-audited)
- Chapters 2-6 with full heading hierarchy and figure placeholders
- Conclusions, References, Glossary
- Page header on body pages, arabic page numbers in the footer

Re-running this script overwrites the existing .docx.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------------------------------------------------------------
# CONFIG  --  edit here when supervisor confirms details
# ----------------------------------------------------------------------
STUDENT_NAME = "Vlad DUMITRU"
COORDINATOR = "Conf. dr. ing. Iuliana MARIN"
TITLE_EN = "Pulsefy: AI-Assisted Ad Creation Platform for Short-Form Social Video"
TITLE_RO = "Pulsefy: Platformă pentru crearea reclamelor asistată de inteligența artificială pentru video-ul scurt din social media"
YEAR = "2026"
SUBMISSION_MONTH = "June 2026"
DEPARTMENT_DIRECTOR = "Prof. dr. ing. Georg DRAGOI"
HEADER_TEXT = (
    f"Diploma Thesis, {STUDENT_NAME}, Faculty of Engineering "
    f"in Foreign Languages, UPB, {YEAR}"
)

OUTPUT_PATH = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


# ----------------------------------------------------------------------
# helpers (identical to v1)
# ----------------------------------------------------------------------
def add_centered(doc, text, size=11, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_blank_lines(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_para(doc, text, size=11, bold=False, italic=False,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_para_with_marker(doc, text, marker, size=11):
    """Body paragraph with a trailing red-italic [marker] for footnote handoff."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    if marker:
        m = p.add_run(f"  [{marker}]")
        m.italic = True
        m.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        m.font.size = Pt(10)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[FIGURE {number}: {caption} — to be drawn]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {number}: {caption}")
    run.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)


def add_table_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[TABLE {number}: {caption} — to be filled]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Table {number}: {caption}")
    run.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)


def add_todo(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"[TODO: {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def page_break(doc):
    doc.add_page_break()


def set_default_font(doc, name="Times New Roman", size=12):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def set_section_margins(section, top=2.5, bottom=2.5, left=2.5, right=2.5):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_uc_table(doc, name, scope, actor, preconditions, postconditions,
                 main_steps, extension=None, table_number=None):
    rows = 6 if extension is None else 7
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    fields = [
        ("Use Case Name:", name),
        ("Scope:", scope),
        ("Actor:", actor),
        ("Preconditions:", preconditions),
        ("Post-conditions:", postconditions),
        ("Main success scenario:", "[Main success scenario — fill in numbered user/system steps]"),
    ]
    if extension:
        fields.append(("Extension:", extension))

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        row.cells[1].text = value

    if table_number:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Table {table_number}: {name} Description")
        run.font.size = Pt(10)
        cap.paragraph_format.space_after = Pt(12)


def add_page_number_field(paragraph):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc_field(paragraph, switches='\\o "1-3" \\h \\z \\u'):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"TOC {switches}"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "Right-click and Update Field in Word to populate."
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    run._r.append(fld_char4)


def add_tof_field(paragraph):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\h \\z \\c "Figure"'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and Update Field in Word to populate."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_text)
    run._r.append(fld_char3)


# ----------------------------------------------------------------------
# build the document
# ----------------------------------------------------------------------
doc = Document()
set_default_font(doc, "Times New Roman", 12)

for level in (1, 2, 3):
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Times New Roman"
    style.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        style.font.size = Pt(18)
    elif level == 2:
        style.font.size = Pt(14)
    else:
        style.font.size = Pt(12)

section = doc.sections[0]
set_section_margins(section)

# ======================================================================
# TITLE PAGE 1  --  English, no logos
# ======================================================================
add_centered(doc, 'UNIVERSITY "POLITEHNICA" OF BUCHAREST', size=12, bold=True)
add_centered(doc, "FACULTY OF ENGINEERING IN FOREIGN LANGUAGES", size=12, bold=True)
add_centered(doc, "COMPUTERS AND INFORMATION TECHNOLOGY - INFORMATION ENGINEERING",
             size=12, bold=True)
add_blank_lines(doc, 8)

add_centered(doc, "DIPLOMA PROJECT", size=28, bold=True)
add_blank_lines(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run("Project coordinator:\n").font.size = Pt(11)
p.add_run(COORDINATOR + "\n").font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Student:\n").font.size = Pt(11)
p.add_run(STUDENT_NAME).font.size = Pt(11)

add_blank_lines(doc, 4)
add_centered(doc, "Bucharest", size=11)
add_centered(doc, YEAR, size=11)
page_break(doc)


# ======================================================================
# TITLE PAGE 2  --  English with logo placeholders
# ======================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Politehnica seal — left]    [FILS logo — right]")
run.italic = True
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

add_centered(doc, 'UNIVERSITY "POLITEHNICA" OF BUCHAREST', size=12, bold=True)
add_centered(doc, "FACULTY OF ENGINEERING", size=12, bold=True)
add_centered(doc, "IN FOREIGN LANGUAGES", size=12, bold=True)
add_centered(doc, "COMPUTERS AND INFORMATION TECHNOLOGY", size=12, bold=True)
add_blank_lines(doc, 8)

add_centered(doc, TITLE_EN, size=22)
add_blank_lines(doc, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Project coordinator                    Student:\n").font.size = Pt(11)
p.add_run(f"{COORDINATOR}                {STUDENT_NAME}").font.size = Pt(11)

add_blank_lines(doc, 2)
add_centered(doc, "Bucharest", size=11)
add_centered(doc, YEAR, size=11)
page_break(doc)


# ======================================================================
# TITLE PAGE 3  --  Romanian with logo placeholders
# ======================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Politehnica seal — left]    [FILS logo — right]")
run.italic = True
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

add_centered(doc, 'UNIVERSITY "POLITEHNICA" OF BUCHAREST', size=12, bold=True)
add_centered(doc, "FACULTY OF ENGINEERING", size=12, bold=True)
add_centered(doc, "IN FOREIGN LANGUAGES", size=12, bold=True)
add_centered(doc, "COMPUTERS AND INFORMATION TECHNOLOGY", size=12, bold=True)
add_blank_lines(doc, 8)

add_centered(doc, TITLE_RO, size=20)
add_blank_lines(doc, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Project coordinator                    Student:\n").font.size = Pt(11)
p.add_run(f"{COORDINATOR}                {STUDENT_NAME}").font.size = Pt(11)

add_blank_lines(doc, 2)
add_centered(doc, "Bucharest", size=11)
add_centered(doc, YEAR, size=11)
page_break(doc)


# ======================================================================
# DIPLOMA PROJECT THEME FORM
# ======================================================================
add_centered(doc, '"POLITEHNICA" UNIVERSITY OF BUCHAREST', size=12, bold=True)
add_centered(doc, "FACULTY OF ENGINEERING IN FOREIGN LANGUAGES", size=12, bold=True)
add_centered(doc, "COMPUTERS AND INFORMATION TECHNOLOGY", size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Approved\n").font.size = Pt(10)
p.add_run("Director of department:\n").font.size = Pt(10)
p.add_run(f"{DEPARTMENT_DIRECTOR}\n").font.size = Pt(10)
p.add_run("[Signature]").italic = True

add_blank_lines(doc, 2)
add_centered(doc, "DIPLOMA PROJECT THEME FOR:", size=14, bold=True)
add_centered(doc, STUDENT_NAME, size=14, bold=True)
add_blank_lines(doc, 1)

theme_fields = [
    ("1. Theme title:",
     [(TITLE_EN, True), (TITLE_RO, True)]),
    ("2. Initial design data:",
     [("The project consists of designing and implementing a web platform for "
       "AI-assisted ad creation, targeting solo creators and small businesses "
       "that need to produce promotional content for short-form social video. "
       "The platform bundles four AI-driven asset-generation subsystems — "
       "scene image generation (Pollinations.ai), background music generation "
       "(Meta MusicGen with a custom procedural fallback), voiceover synthesis "
       "(gTTS), and a video assembly pipeline (FFmpeg overlay) — into a single "
       "web application with role-based access for visitors, authenticated "
       "users, and administrators. Output formats target the dominant "
       "short-form social platforms (TikTok 9:16, Instagram Reels 9:16, "
       "Instagram square 1:1, YouTube Shorts 9:16, YouTube 16:9).", True)]),
    ("3. Student contribution:",
     [("Bibliographical research", True),
      ("Project analysis", True),
      ("AI generation pipeline implementation", True),
      ("Web application development", True)]),
    ("4. Compulsory graphical material:",
     [("Block scheme, functioning diagram, graphs", True)]),
    ("5. The paper is based on the knowledge obtained at the following study courses:",
     [("Programming Languages, Databases, Object Oriented Programming, Software "
       "Development Methods, Web Applications Development, Artificial Intelligence", True)]),
    ("6. Paper development environment:",
     [("Visual Studio Code, DataGrip, Postman, Microsoft Word, draw.io, Mermaid, "
       "Canva", True)]),
    ("7. The paper serves as:",
     [("Didactic purposes", True)]),
    ("8. Paper preparation date:",
     [(SUBMISSION_MONTH, True)]),
]

for label, items in theme_fields:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    for text, italic in items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        run = para.add_run(text)
        run.italic = italic
        run.font.size = Pt(11)

add_blank_lines(doc, 2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run("Project coordinator                                                          "
         "Student:\n").bold = True
p.add_run(f"{COORDINATOR}                                                            "
          f"{STUDENT_NAME}\n")
p.add_run("[Signature]").italic = True
p.add_run("                                                                                                         "
          "[Signature]").italic = True
page_break(doc)


# ======================================================================
# ACADEMIC HONESTY STATEMENT
# ======================================================================
add_centered(doc, "Academic Honesty Statement", size=18, bold=True)
add_blank_lines(doc, 2)

honesty_paragraphs = [
    f'I, {STUDENT_NAME}, hereby declare that the work with the title '
    f'"{TITLE_EN}", to be openly defended in front of the diploma theses '
    f'examination commission at the Faculty of Engineering in Foreign '
    f'Languages, University "Politehnica" of Bucharest, as partial requirement '
    f'for obtaining the title of Engineer is the result of my own work, based '
    f'on my work.',

    "The thesis, simulations, experiments and measurements that are presented "
    "are made entirely by me under the guidance of the scientific adviser, "
    "without the implication of persons that are not cited by name and "
    "contribution in the Acknowledgements part.",

    "The thesis has never been presented to a higher education institution or "
    "research board in the country or abroad.",

    "All the information used, including the Internet, is obtained from sources "
    "that were cited and indicated in the notes and in the bibliography, "
    "according to ethical standards. I understand that plagiarism is an offense "
    "and is punishable under law.",

    "The results from the simulations, experiments and measurements are "
    "genuine. I understand that the falsification of data and results "
    "constitutes fraud and is punished according to regulations.",
]
for para in honesty_paragraphs:
    add_para(doc, para, size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=12)

add_blank_lines(doc, 4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run(f"{STUDENT_NAME}                                                              "
          f"[Date]\n")
p.add_run("[Signature]").italic = True
page_break(doc)


# ======================================================================
# TABLE OF CONTENTS
# ======================================================================
add_centered(doc, "Table of Contents", size=18, bold=True)
add_blank_lines(doc, 1)
toc_para = doc.add_paragraph()
add_toc_field(toc_para)
page_break(doc)


# ======================================================================
# TABLE OF FIGURES
# ======================================================================
add_centered(doc, "Table of Figures", size=18, bold=True)
add_blank_lines(doc, 1)
tof_para = doc.add_paragraph()
add_tof_field(tof_para)
page_break(doc)


# ======================================================================
# BODY  --  separate section for page header + page numbering
# ======================================================================
new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
new_section.different_first_page_header_footer = False
set_section_margins(new_section)

new_section.header.is_linked_to_previous = False
new_section.footer.is_linked_to_previous = False

header = new_section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run = header_para.add_run(HEADER_TEXT)
header_run.font.size = Pt(10)
pPr = header_para._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "auto")
pBdr.append(bottom)
pPr.append(pBdr)

footer = new_section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number_field(footer_para)


# ======================================================================
# CHAPTER 1  --  INTRODUCTION  (full prose, 5 subsections)
# ======================================================================
add_heading_styled(doc, "1 Introduction", level=1)

# ---- 1.1 ----
add_heading_styled(doc, "1.1 The shift to social-video advertising", level=2)
add_para_with_marker(doc,
    "Short-form vertical video has become the dominant format for "
    "social-media advertising. TikTok, Instagram Reels, and YouTube Shorts "
    "each report user audiences in the hundreds of millions, with watch "
    "sessions composed of dozens of clips per sitting. Brands and creators "
    "that need to reach those audiences are concentrating their advertising "
    "effort on this format because it matches where viewer attention has "
    "moved.",
    "footnote F1: TikTok / Meta active user figures and short-form ad "
    "engagement — see THESIS_RESEARCH_BANK.md Section A; verify and add "
    "Word footnote here")
add_para_with_marker(doc,
    "The advertising spend has followed the attention. Global digital "
    "advertising revenue continues to grow year on year, with social-media "
    "platforms taking a steadily larger share, and short-form video has been "
    "the fastest-growing single format on those platforms.",
    "footnote F2: digital ad spend trend — cite eMarketer / WARC / IFPI for "
    "comparable music-industry digital share if helpful; add Word footnote")
add_para(doc,
    "Three properties of short-form social video distinguish it from older "
    "online ad formats. First, the 9:16 vertical aspect ratio matches phone "
    "screens directly, without letterboxing or rotation. Second, the "
    "algorithmic feed delivers content from accounts the viewer has never "
    "followed, which means an ad on a small account can still reach a large "
    "audience if the creative performs in the first few seconds. Third, the "
    "format is audio-on by default, with most users watching with sound; "
    "this makes voiceover, music, and sound design first-class production "
    "decisions rather than afterthoughts.")
add_para(doc,
    "The fourth property is cadence. The recommendation algorithms on TikTok "
    "and Reels reward accounts that post consistently — daily or several "
    "times per day. A small business that wants steady reach has to produce "
    "new ad content at a frequency that previous-generation production "
    "pipelines were never built for. A creator who can ship one polished ad "
    "in a working day is at a disadvantage against a competitor who can ship "
    "three using AI assistance.")
add_para(doc,
    "These properties together define the production environment Pulsefy "
    "targets: vertical format, audio-rich, frequent, and produced by people "
    "who do not have the budget for a traditional creative agency.")
add_figure_placeholder(doc, 1, "Pulsefy ad creation pipeline overview")

# ---- 1.2 ----
add_heading_styled(doc, "1.2 Who creates ads now", level=2)
add_para(doc,
    "For most of advertising's history, producing a campaign ad required a "
    "creative agency. The agency briefed a copywriter, hired a director and "
    "crew, booked a vocal artist for voiceover, contracted a music house for "
    "the soundtrack, and ran the rough cut through a post-production team. "
    "The end product was professional. The cost was several thousand euros "
    "per campaign and the turnaround was weeks.")
add_para(doc,
    "Short-form social advertising has eroded this model from below. Brands "
    "now also commission ads from individual creators on retainer, in-house "
    "marketing teams produce campaigns directly, and a substantial share of "
    "paid promotion on TikTok and Reels is created by the founder or owner "
    "of the business being advertised. The same person designing the product "
    "writes the ad copy, records the voiceover, picks the soundtrack, and "
    "edits the short video.")
add_para(doc,
    "This shift introduces a real skill spread problem. A retail entrepreneur "
    "who runs a specialty food shop knows the product but is not a video "
    "editor. A side-hustle creator selling a digital tool can write good "
    "copy but cannot produce broadcast-quality voiceover. A small marketing "
    "team in a startup of five people may be excellent at strategy but have "
    "no in-house designer for product photography. Each of these creators is "
    "producing ads anyway, because their reach depends on it.")
add_para(doc,
    "What this segment of users has in common is that they need "
    "professional-looking ad output without the time or budget to develop "
    "professional production skills, and without the budget to hire a "
    "creative agency on retainer. They need the production tools to "
    "compensate for the missing in-house specialists.")
add_para(doc,
    "Pulsefy targets this segment specifically. The application is designed "
    "for solo creators, small business owners, and in-house marketers who "
    "need to produce frequent short-form ads without learning a separate "
    "professional tool for each stage of production.")

# ---- 1.3 ----
add_heading_styled(doc, "1.3 The cost gap and tool fragmentation problem", level=2)
add_para_with_marker(doc,
    "The cost gap between what small advertisers need and what is available "
    "to them has two layers. The first is the absolute cost of "
    "creative-agency work, which still ranges from one to ten thousand euros "
    "per campaign for the kind of small project a side-hustle creator might "
    "commission. For a retailer with a few hundred euros in monthly ad "
    "budget, this is unaffordable as a recurring expense.",
    "footnote F3: agency price range — cite a Romanian or European agency "
    "rate card; add Word footnote here")
add_para_with_marker(doc,
    "The second layer is the subscription cost of professional production "
    "software. Adobe Creative Cloud, the industry-standard package for "
    "design, image editing, and video, costs around sixty euros per month. "
    "Final Cut Pro, DaVinci Resolve Studio, Logic Pro, and similar "
    "professional tools each carry one-time or subscription fees, on top of "
    "which most have steep learning curves measured in months rather than "
    "days. The total stack is unattractive for someone producing a handful "
    "of ads per month.",
    "footnote F4: Adobe Creative Cloud monthly price — cite official Adobe "
    "pricing page; add Word footnote here")
add_para(doc,
    "The third layer is the do-it-yourself path that most small advertisers "
    "actually take, which involves stitching together several free or "
    "low-cost specialised tools. Design happens in Canva or Figma. Video "
    "editing happens in CapCut or InShot. Background music is generated in "
    "Suno or licensed from a stock-music site. Voiceover is generated in "
    "ElevenLabs or recorded on a phone. Final assembly happens back in "
    "CapCut. Each tool has its own subscription, its own export format, its "
    "own learning curve, and its own user interface. The user pays multiple "
    "subscriptions, switches contexts repeatedly to make a single ad, and "
    "has to manually convert files between tools when formats do not match.")
add_para(doc,
    "In practice, producing a single short-form ad through this fragmented "
    "pipeline takes several hours of focused work. The actual creative "
    "input — what the ad should communicate — is a small share of that "
    "time. The bulk goes to file conversion, format adjustment, asset "
    "transfer between applications, and re-learning interfaces that were "
    "used a few weeks ago and partially forgotten in the meantime.")
add_para(doc,
    "This fragmentation is the core problem Pulsefy is built to address. "
    "Each tool in the existing landscape is excellent at one stage of the "
    "production pipeline, but none brings the stages together under one "
    "account, one subscription, and one consistent interface.")
add_figure_placeholder(doc, 2,
    "Manual versus AI-assisted ad creation timelines")

# ---- 1.4 ----
add_heading_styled(doc, "1.4 Advantages of AI-assisted ad creation", level=2)
add_para(doc,
    "AI-assisted ad creation tools offer concrete advantages over both the "
    "agency model and the fragmented do-it-yourself path. The first "
    "advantage is asset generation at scale. Image-generation models can "
    "produce a product hero shot, a lifestyle context image, a close-up "
    "detail, and a brand-moment image from the same product description, in "
    "less time than it takes to brief a photographer. Music-generation "
    "models can produce royalty-free background tracks at any tempo and mood "
    "without licensing negotiations. Text-to-speech models can produce "
    "voiceover in dozens of languages without scheduling a vocal artist.")
add_para(doc,
    "The second advantage is multi-language reach. A small business "
    "advertising in Romania can use the same campaign script in English, "
    "French, German, and Spanish, with the voiceover regenerated in each "
    "language at no additional production cost. This is structurally "
    "impossible to do cheaply in the agency model, where each language "
    "version requires a separate vocal artist booking.")
add_para(doc,
    "The third advantage is variant generation. Short-form social platforms "
    "reward A/B testing of multiple ad variants — different opening hooks, "
    "different soundtracks, different visual treatments — and the algorithms "
    "pick the winner. AI generation makes producing five variants of the "
    "same ad almost as cheap as producing one, which makes systematic A/B "
    "testing accessible to creators who could previously only afford to "
    "produce a single version.")
add_para(doc,
    "The fourth advantage is turnaround speed. The cadence demanded by "
    "TikTok and Reels algorithms — daily or near-daily output — is "
    "achievable only with significant tool acceleration. AI generation "
    "reduces the time from brief to finished ad from hours to minutes for "
    "the asset-creation stages, leaving the human creator to focus on the "
    "brief itself and on quality control.")
add_para(doc,
    "The trade-off, which Pulsefy acknowledges and partially addresses, is "
    "that AI-generated assets remain lower in polish than what a senior "
    "creative team would produce, and that the abundance of generated "
    "content can lead to homogeneity if the creator does not exert taste "
    "over the output. Both limits are real and discussed in the validation "
    "chapter.")

# ---- 1.5 ----
add_heading_styled(doc, "1.5 Structure of the paper", level=2)
add_para(doc,
    "The aim of this paper is to present Pulsefy, a web application for "
    "AI-assisted ad creation, and to document the design, implementation, "
    "and validation of its components.")
add_para(doc,
    "The paper is organised in six chapters, plus the conclusions and "
    "supporting back matter:")
chapters_list = [
    "Chapter 1 introduces the motivation behind the project and the context in which it sits.",
    "Chapter 2 reviews scientific work on generative media for advertising and "
        "compares four commercial ad-creation platforms with an overlapping focus.",
    "Chapter 3 specifies the functional and non-functional requirements of the application.",
    "Chapter 4 describes the architectural and detailed design, including class diagrams, "
        "sequence diagrams, and database structure.",
    "Chapter 5 documents the implementation, including the technologies used, "
        "the AI generation subsystems, and a walkthrough of the web interface from an "
        "advertiser's perspective.",
    "Chapter 6 presents the validation methodology and results.",
]
for item in chapters_list:
    p = doc.add_paragraph(item, style="List Bullet")
add_para(doc,
    "The diagrams in this paper were created using Mermaid "
    "(https://mermaid.live) and exported as PNG. The decorative figures "
    "were created using Canva (https://www.canva.com/). The application "
    "source code was written using Visual Studio Code, and the database "
    "scripts in DataGrip.")
page_break(doc)


# ======================================================================
# CHAPTER 2  --  STATE OF THE ART
# ======================================================================
add_heading_styled(doc, "2 State of the art", level=1)
add_todo(doc, "intro paragraph framing scientific vs commercial split for "
              "ad creation")

add_heading_styled(doc, "2.1 Scientific state of the art", level=2)
add_todo(doc, "intro paragraph anchoring the project's three ML domains: "
              "text-to-image, text-to-music, text-to-speech")

add_heading_styled(doc, "2.1.1 Text-to-image generation for advertising", level=3)
add_todo(doc, "prose: latent diffusion (Stable Diffusion, Flux), "
              "Pollinations.ai usage; cite original DDPM and SD papers, plus "
              "ad-photography use-case literature")
add_figure_placeholder(doc, 3,
    "Existing text-to-image methods for ad scene generation")

add_heading_styled(doc, "2.1.2 Generative audio models for ad soundtracks", level=3)
add_todo(doc, "prose: MusicGen, MusicLM, AudioLDM, Stable Audio. Cite Copet "
              "et al. 2023 (MusicGen) as the central reference since Pulsefy "
              "uses it directly")

add_heading_styled(doc, "2.1.3 Text-to-speech for ad voiceover", level=3)
add_todo(doc, "prose: gTTS as a wrapper around Google Translate TTS, "
              "comparison with Tacotron / VITS / ElevenLabs neural TTS, "
              "trade-offs (free vs quality)")

add_heading_styled(doc, "2.2 Commercial state of the art", level=2)
add_todo(doc, "Open with concrete industry case-study hook plus bulleted "
              "result stats. Then bulleted list of the four tools to be "
              "analyzed.")

add_heading_styled(doc, "2.2.1 Canva", level=3)
add_todo(doc, "identity sentence; design + video workflow with Magic AI "
              "features; quantitative outcomes (paid users, ad creator base); "
              "API/pricing")
add_figure_placeholder(doc, 4, "Canva ad creation workflow")

add_heading_styled(doc, "2.2.2 CapCut", level=3)
add_todo(doc, "identity sentence (ByteDance, free, TikTok-native); video "
              "editing with AI templates and effects; outcomes; pricing")
add_figure_placeholder(doc, 5, "CapCut AI templates and effects")

add_heading_styled(doc, "2.2.3 AdCreative.ai", level=3)
add_todo(doc, "identity sentence; AI-driven static and video ad generation; "
              "outcomes; pricing tiers (premium-leaning)")
add_figure_placeholder(doc, 6, "AdCreative.ai generation flow")

add_heading_styled(doc, "2.2.4 InVideo AI", level=3)
add_todo(doc, "identity sentence; AI video generation for ads from text "
              "prompts; outcomes; pricing")
add_figure_placeholder(doc, 7, "InVideo AI generation flow")

add_heading_styled(doc, "2.3 Summary and comparison of existing approaches", level=2)
add_todo(doc, "summary prose pointing at the comparison table")
add_table_placeholder(doc, 1,
    "Comparison of commercial AI ad creation platforms")
page_break(doc)


# ======================================================================
# CHAPTER 3  --  REQUIREMENTS ANALYSIS
# ======================================================================
add_heading_styled(doc, "3 Requirements analysis", level=1)

add_heading_styled(doc, "3.1 Functional requirements", level=2)

add_heading_styled(doc, "3.1.1 Actors and agents", level=3)
add_todo(doc, "two short prose paragraphs: Visitor + Authenticated User as "
              "actors (free vs premium tier handled as a precondition on "
              "premium-only use cases); Administrator as agent")

add_heading_styled(doc, "3.1.2 Software use case diagram", level=3)
add_figure_placeholder(doc, 8, "Software use case diagram")

add_heading_styled(doc, "3.1.3 Use case functional requirements", level=3)
add_todo(doc, "list every UC as 3.1.3.X with lettered system-does-X bullet "
              "list. If using priority weights, add explanatory paragraph at "
              "the start.")

add_heading_styled(doc, "3.1.4 Use case descriptions and system sequence diagrams",
                   level=3)
ucs = [
    ("UC1 Register", "Visitor",
     "The visitor accesses the registration form.",
     "The user account is created and the visitor is logged in."),
    ("UC2 Log in", "Visitor",
     "The user has a registered account.",
     "The user is logged in and the dashboard is displayed."),
    ("UC3 Create new ad project", "Authenticated User",
     "The user is logged in.",
     "A new ad project is created and shown in the user's project list."),
    ("UC4 Generate ad scene images", "Authenticated User",
     "The user has an active ad project and a product description.",
     "The requested number of scene images is generated and attached to the project."),
    ("UC5 Generate background music", "Authenticated User",
     "The user has an active ad project and a music brief.",
     "The generated track is attached to the project; failures are surfaced."),
    ("UC6 Browse licensed music catalog", "Authenticated User",
     "The user is logged in.",
     "The user selects a licensed track and attaches it to the project."),
    ("UC7 Generate voiceover narration", "Authenticated User",
     "The user has an ad project and a script.",
     "The voiceover audio is generated and attached to the project."),
    ("UC8 Upload product video and overlay audio", "Authenticated User",
     "The user has a video file and a chosen audio source.",
     "The processed video is stored and a download link is shown."),
    ("UC9 Assemble and export ad", "Authenticated User",
     "The user has an ad project with at least one image, an audio source, "
     "and an aspect-ratio selection.",
     "A finished ad video is rendered in the requested format and made available for download."),
    ("UC10 Manage subscription tier", "Authenticated User",
     "The user is logged in.",
     "The user's subscription tier is updated and persisted."),
    ("UC11 Manage users and content", "Administrator",
     "The administrator is logged in.",
     "The user roster or moderation queue is updated."),
]
for i, (uc_name, actor, pre, post) in enumerate(ucs, start=1):
    add_heading_styled(doc, f"3.1.4.{i} {uc_name}", level=3)
    add_uc_table(
        doc,
        name=uc_name,
        scope="System",
        actor=actor,
        preconditions=pre,
        postconditions=post,
        main_steps=None,
        extension=None,
        table_number=i + 1,
    )
    add_todo(doc, f"fill in main success scenario steps; add Extension row "
                  f"if alt flow exists")
    add_figure_placeholder(doc, 8 + i,
                           f"{uc_name} system sequence diagram")

add_heading_styled(doc, "3.1.5 Activity diagrams", level=3)
add_todo(doc, "build activity diagrams for the 2-3 most branching use cases")
add_figure_placeholder(doc, 20,
    "Generate ad scene images activity diagram")
add_figure_placeholder(doc, 21,
    "Assemble and export ad activity diagram")

add_heading_styled(doc, "3.1.6 Operation contracts", level=3)
add_todo(doc, "3-4 contract tables for: GenerateAdImages, GenerateMusic, "
              "GenerateVoiceover, AssembleAndExportAd. Same field labels as "
              "IARCA: Operation Name (signature) / Use Case Name / "
              "Preconditions / Post-conditions.")

add_heading_styled(doc, "3.2 Non-functional requirements", level=2)
nfr_list = ["Performance", "Security", "Usability", "Reliability",
            "Scalability", "Maintainability", "Portability"]
for i, nfr in enumerate(nfr_list, start=1):
    add_heading_styled(doc, f"3.2.{i} {nfr}", level=3)
    add_todo(doc, f"2-4 sentences defending the {nfr.lower()} target for an "
                  "ad-creation web platform")
page_break(doc)


# ======================================================================
# CHAPTER 4  --  DESIGN
# ======================================================================
add_heading_styled(doc, "4 Design", level=1)

add_heading_styled(doc, "4.1 Architecture", level=2)
add_heading_styled(doc, "4.1.1 Architecture choice", level=3)
add_todo(doc, "3 short paragraphs: definition of client-server with separate "
              "Python ML services for the four AI subsystems, drawback, then "
              "'fits because…' rationale.")
add_figure_placeholder(doc, 22, "Pulsefy architecture block diagram")

add_heading_styled(doc, "4.1.2 Package diagrams", level=3)
add_figure_placeholder(doc, 23, "pulsefy.server package diagram")
add_figure_placeholder(doc, 24, "pulsefy.client package diagram")

add_heading_styled(doc, "4.1.3 Deployment", level=3)
add_todo(doc, "describe local dev deployment (Vite :8080, Node :4000, "
              "Postgres). Cloud target if planned.")
add_figure_placeholder(doc, 25, "Deployment configuration screenshot")

add_heading_styled(doc, "4.2 Detailed design", level=2)
add_heading_styled(doc, "4.2.1 Design class diagrams and design sequence diagrams",
                   level=3)
fig_idx = 26
for i, (uc_name, *_) in enumerate(ucs, start=1):
    add_heading_styled(doc, f"4.2.1.{i} {uc_name}", level=3)
    add_figure_placeholder(doc, fig_idx,
                           f"{uc_name} design class diagram")
    fig_idx += 1
    add_figure_placeholder(doc, fig_idx,
                           f"{uc_name} design sequence diagram")
    fig_idx += 1

add_heading_styled(doc, "4.2.2 Database structure", level=3)
add_todo(doc, "open with the IARCA-style 'three points of view' framing: "
              "AI subsystem schema, web app at deployment, web app at "
              "runtime with example rows")
add_figure_placeholder(doc, fig_idx, "Pulsefy AI subsystem ERD")
fig_idx += 1
add_figure_placeholder(doc, fig_idx, "Pulsefy web app ERD at deployment")
fig_idx += 1
add_figure_placeholder(doc, fig_idx,
    "Pulsefy web app ERD at runtime with example rows")
fig_idx += 1
page_break(doc)


# ======================================================================
# CHAPTER 5  --  IMPLEMENTATION
# ======================================================================
add_heading_styled(doc, "5 Implementation", level=1)

add_heading_styled(doc, "5.1 Technologies", level=2)
add_heading_styled(doc, "5.1.1 Development", level=3)
add_todo(doc, "three running paragraphs: front-end stack (React, TS, Vite, "
              "Tailwind, shadcn), back-end stack (Node 18, Express 4, "
              "raw pg, JWT, bcrypt, Zod), AI/ML services (Python 3.11 venv "
              "shared by MusicGen, gTTS; Pollinations HTTP API)")
add_figure_placeholder(doc, fig_idx, "Technologies used (logo ribbon)")
fig_idx += 1

add_heading_styled(doc, "5.1.2 External services", level=3)
add_todo(doc, "Pollinations.ai for scene image generation, Jamendo Catalog "
              "API for licensed music, Hugging Face model hub for "
              "audiocraft/MusicGen — describe endpoints, quotas, retry "
              "strategy")

add_heading_styled(doc, "5.2 AI generation subsystems", level=2)
add_todo(doc, "intro sentence + 4-column matrix table summarising the four "
              "AI subsystems and their public functions")
add_table_placeholder(doc, 2,
    "AI generation subsystems and their public functions")

add_heading_styled(doc, "5.2.1 Music generation (MusicGen and Pulsefy AI)", level=3)
add_todo(doc, "MusicGen subprocess invocation, Python script flow, output "
              "WAV format, Pulsefy AI procedural fallback for free tier; "
              "include pseudocode block for the spawn + stream loop")
add_figure_placeholder(doc, fig_idx,
    "Pseudocode: music generation subsystem")
fig_idx += 1

add_heading_styled(doc, "5.2.2 Voiceover generation (gTTS)", level=3)
add_todo(doc, "gTTS Python script, MP3 output, ffmpeg conversion to "
              "44.1 kHz mono 16-bit PCM WAV, language list, error handling. "
              "Include pseudocode block for the gTTS + ffmpeg pipeline")

add_heading_styled(doc, "5.2.3 Scene image generation (Pollinations.ai)", level=3)
add_todo(doc, "Pollinations HTTP endpoint, prompt assembly with "
              "scene-angle and style modifiers, retry/backoff on 429s, "
              "aspect ratio handling")
add_figure_placeholder(doc, fig_idx,
    "Pseudocode: scene image generation subsystem")
fig_idx += 1

add_heading_styled(doc, "5.2.4 Video assembly (FFmpeg overlay)", level=3)
add_todo(doc, "video upload validation, ffmpeg command for overlaying "
              "generated audio onto user-provided video, aspect ratio "
              "handling for export. Include the full ffmpeg command line "
              "as a code listing")
add_figure_placeholder(doc, fig_idx,
    "Sample ffmpeg command for video and audio assembly")
fig_idx += 1

add_heading_styled(doc, "5.2.5 Configuration", level=3)
add_todo(doc, "describe the .env structure with bullets: MUSICGEN_MODEL, "
              "MUSICGEN_DURATION_SEC, MUSICGEN_DEVICE, GTTS_LANGUAGE, "
              "POLLINATIONS_BASE_URL, JAMENDO_CLIENT_ID, DATABASE_URL, JWT_SECRET")
add_figure_placeholder(doc, fig_idx,
    "App settings structure (.env example)")
fig_idx += 1

add_heading_styled(doc, "5.3 Web application walkthrough", level=2)

add_heading_styled(doc, "5.3.1 Walkthrough for an advertiser building a campaign",
                   level=3)
advertiser_flows = [
    "Login and signup",
    "Dashboard with active ad projects",
    "Create new ad project",
    "Enter product brief (name, description, target platform)",
    "Generate ad scene images and pick from variants",
    "Generate background music or pick from licensed catalog",
    "Generate voiceover narration in target language",
    "Upload product video and overlay generated audio",
    "Assemble and preview the final ad",
    "Export ad in target aspect ratio (9:16, 1:1, or 16:9)",
]
for flow in advertiser_flows:
    add_figure_placeholder(doc, fig_idx, f"Pulsefy – {flow}")
    fig_idx += 1
    add_todo(doc, f"3 paragraphs describing the {flow.lower()} step")

add_heading_styled(doc, "5.3.2 Administrator panel walkthrough", level=3)
admin_flows = [
    "User management",
    "Subscription tier moderation",
    "Generated content moderation",
]
for flow in admin_flows:
    add_figure_placeholder(doc, fig_idx, f"Pulsefy – {flow}")
    fig_idx += 1
    add_todo(doc, f"2 paragraphs describing the {flow.lower()} step")
page_break(doc)


# ======================================================================
# CHAPTER 6  --  VALIDATION
# ======================================================================
add_heading_styled(doc, "6 Validation", level=1)
add_todo(doc,
    "validation chapter combining: per-subsystem latency benchmarks (real "
    "measurements with documented protocol), small N=15 ad-quality pilot "
    "study (real measurements), system capacity metrics. DO NOT INVENT "
    "NUMBERS. Each metric needs a documented formula or measurement protocol.")
add_figure_placeholder(doc, fig_idx,
    "Pulsefy validation summary infographic")
fig_idx += 1
page_break(doc)


# ======================================================================
# CONCLUSIONS
# ======================================================================
add_heading_styled(doc, "Conclusions", level=1)
add_todo(doc,
    "five-block structure: (1) recap thesis claim — Pulsefy unifies "
    "ad-creation asset generation in one product, (2) bullet list of design "
    "constraints surfaced during analysis, (3) one personal first-person "
    "sentence, (4) limitations + future work (image gen route wiring, "
    "production deployment, expanded language coverage), (5) short outro.")
page_break(doc)


# ======================================================================
# REFERENCES
# ======================================================================
add_heading_styled(doc, "References", level=1)
add_todo(doc, "25-35 numbered [N] entries with hanging indent. APA-hybrid "
              "format. See docs/THESIS_RESEARCH_BANK.md for verified "
              "citation skeletons; add ad-creation competitor citations "
              "(Canva, CapCut, AdCreative.ai, InVideo AI) once researched.")
page_break(doc)


# ======================================================================
# GLOSSARY
# ======================================================================
add_heading_styled(doc, "Glossary", level=1)
add_para(doc,
    "The terms below are abbreviations and technical concepts referenced in "
    "this paper. Proper nouns (such as MusicGen, Jamendo, Pollinations.ai) "
    "are not listed here.",
    space_after=12,
)

glossary_terms = [
    ("API", "Application Programming Interface"),
    ("BPM", "Beats Per Minute"),
    ("CC", "Creative Commons (license family used by the Jamendo catalog)"),
    ("CDN", "Content Delivery Network"),
    ("CLI", "Command Line Interface"),
    ("CPM", "Cost Per Mille (cost per one thousand ad impressions)"),
    ("CPU", "Central Processing Unit"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("CSS", "Cascading Style Sheets"),
    ("CSV", "Comma-Separated Values"),
    ("CTA", "Call to Action"),
    ("CTR", "Click-Through Rate"),
    ("DAW", "Digital Audio Workstation"),
    ("DSP", "Digital Service Provider (or Digital Signal Processor)"),
    ("ERD", "Entity-Relationship Diagram"),
    ("FK", "Foreign Key"),
    ("GB", "Gigabyte"),
    ("GPU", "Graphics Processing Unit"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("HTTPS", "HyperText Transfer Protocol Secure"),
    ("IDE", "Integrated Development Environment"),
    ("JS", "JavaScript"),
    ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"),
    ("KPI", "Key Performance Indicator"),
    ("LLM", "Large Language Model"),
    ("MAU", "Monthly Active Users"),
    ("MB", "Megabyte"),
    ("ML", "Machine Learning"),
    ("M:N", "Many-to-Many (relational cardinality)"),
    ("MoSCoW", "Must, Should, Could, Won't (requirement priority scale)"),
    ("MP3", "MPEG Audio Layer III"),
    ("MVC", "Model-View-Controller"),
    ("NLP", "Natural Language Processing"),
    ("npm", "Node Package Manager"),
    ("OOP", "Object-Oriented Programming"),
    ("ORM", "Object-Relational Mapping"),
    ("OS", "Operating System"),
    ("PCM", "Pulse Code Modulation"),
    ("PDF", "Portable Document Format"),
    ("PK", "Primary Key"),
    ("PNG", "Portable Network Graphics"),
    ("REST", "Representational State Transfer"),
    ("ROAS", "Return On Ad Spend"),
    ("SaaS", "Software as a Service"),
    ("SDK", "Software Development Kit"),
    ("SPA", "Single-Page Application"),
    ("SQL", "Structured Query Language"),
    ("TS", "TypeScript"),
    ("TTS", "Text-to-Speech"),
    ("UC", "Use Case"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
    ("UX", "User Experience"),
    ("venv", "Virtual Environment (Python)"),
    ("WAV", "Waveform Audio File Format"),
    ("XML", "eXtensible Markup Language"),
]

glossary_table = doc.add_table(rows=len(glossary_terms), cols=2)
glossary_table.autofit = False
for i, (term, definition) in enumerate(glossary_terms):
    cell_term = glossary_table.rows[i].cells[0]
    cell_term.text = term
    for p in cell_term.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)
    cell_def = glossary_table.rows[i].cells[1]
    cell_def.text = definition
    for p in cell_def.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)

for row in glossary_table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(13)


# ======================================================================
# save
# ======================================================================
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
