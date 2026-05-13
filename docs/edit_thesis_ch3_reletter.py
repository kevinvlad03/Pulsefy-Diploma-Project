"""Re-letter UC4, UC5, UC7 items in 3.1.3 after the quota-line removals."""

import re
from docx import Document
from docx.shared import Pt

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def reletter_uc_items(doc, heading_signature):
    """Walk through lettered items between this heading and the next sub-sub-heading
    (text starting with "3.1.3." or styled as a Heading).

    Re-letter them a), b), c) ... in order, preserving the remainder of each line."""
    import re as _re
    letters = "abcdefghij"
    paras = list(doc.paragraphs)

    start = None
    for i, p in enumerate(paras):
        if heading_signature in p.text:
            start = i
            break
    if start is None:
        return 0

    # End at the next paragraph whose text begins with "3.1.3." (next sub-sub-section)
    # or is a real Heading-styled paragraph.
    next_marker = _re.compile(r"^3\.1\.[3-9](\.|\s|$)")
    end = len(paras)
    for i in range(start + 1, len(paras)):
        if paras[i].style.name.startswith("Heading"):
            end = i
            break
        if next_marker.match(paras[i].text.strip()):
            end = i
            break

    idx = 0
    pattern = re.compile(r"^([a-z])\)\s+(.*)$")
    for i in range(start + 1, end):
        p = paras[i]
        text = p.text.strip()
        if not text:
            continue
        m = pattern.match(text)
        if m:
            new_text = f"{letters[idx]}) {m.group(2)}"
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            new_run = p.add_run(new_text)
            new_run.font.size = Pt(11)
            idx += 1
    return idx


doc = Document(DOCX)
c4 = reletter_uc_items(doc, "3.1.3.4 Generate ad scene images")
c5 = reletter_uc_items(doc, "3.1.3.5 Generate background music")
c7 = reletter_uc_items(doc, "3.1.3.7 Generate voiceover narration")
print(f"UC4: {c4} items re-lettered")
print(f"UC5: {c5} items re-lettered")
print(f"UC7: {c7} items re-lettered")
doc.save(DOCX)
print(f"Saved: {DOCX}")
