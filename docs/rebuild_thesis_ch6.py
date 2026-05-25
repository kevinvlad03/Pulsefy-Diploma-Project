"""
Rebuild Chapter 6 (Validation) of Pulsefy_Diploma_Thesis.docx.

Replaces everything between the existing '6 Validation' heading and the
'Conclusions' heading with the real measured benchmark.

Numbers come from the three end-to-end runs Vlad measured on his
Mac M2 Pro under Instagram Reel settings (12-second music, balanced
MusicGen, English gTTS, three 9:16 Pollinations scenes per ad).

Run from the worktree root:
    python3 docs/rebuild_thesis_ch6.py
"""

from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def add_heading_styled(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = "Times New Roman"
    return h


def add_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_figure_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[FIGURE {number}: {caption} — to be inserted]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {number}: {caption}")
    run.font.size = Pt(10)


def add_results_table(doc, table_number, caption):
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"

    headers = ["Stage", "Case 1: Coffee", "Case 2: Sneaker", "Case 3: Skincare", "Average"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    data_rows = [
        ("Pulsefy AI music (LSTM, 12 s)", "7 s", "3 s", "3 s", "~4 s"),
        ("MusicGen balanced (12 s)", "7 min 00 s", "5 min 30 s", "5 min 45 s", "~6 min 05 s"),
        ("gTTS voiceover (English)", "2 s", "2 s", "2 s", "~2 s"),
        ("Video assembly (3 scenes + audio overlay)", "3 min 00 s", "3 min 00 s", "3 min 15 s", "~3 min 05 s"),
    ]
    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            if j == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Table {table_number}: {caption}")
    run.font.size = Pt(10)


temp = Document()
add_heading_styled(temp, "6 Validation", level=1)

add_para(temp,
    "This chapter reports the end-to-end latency of the Pulsefy ad-creation "
    "pipeline measured under realistic Instagram Reel conditions. The "
    "benchmark covers three ad-creation scenarios, each one a full pipeline "
    "run using the default settings a typical advertiser would select. The "
    "aim is not to compare Pulsefy against other platforms or to evaluate "
    "the quality of the generated content; it is to verify that the "
    "platform produces a finished ad in minutes rather than hours, as "
    "claimed in Chapter 1."
)
add_para(temp,
    "All benchmarks were run on a 2023 Apple Mac M2 Pro (12-core CPU, "
    "19-core GPU, 16 GB unified memory) running macOS Sonoma. The Express "
    "server, the React frontend, the Python virtual environment, and the "
    "PostgreSQL database all ran locally on the same machine. The MusicGen "
    "subprocess used the Apple Metal (MPS) backend. The three scenarios "
    "shared a common configuration: 12-second music duration, balanced "
    "MusicGen quality, English gTTS without slow-talk, and three "
    "Pollinations.ai-generated 9:16 portrait scene images per ad. The "
    "scenarios differed only in their product theme, music prompt, "
    "voiceover script, and image prompt (a coffee brand, a sneaker drop, "
    "and a skincare product). Latency was measured with a stopwatch from "
    "the moment each form was submitted to the moment the resulting asset "
    "appeared in the user interface."
)
add_para(temp,
    "Table 3 reports the four measured stages across the three scenarios. "
    "The Pulsefy AI music engine produced the requested track in single-"
    "digit seconds on every run, with the longer Case 1 measurement (7 "
    "seconds) reflecting the first cold-start invocation of the LSTM "
    "checkpoint. MusicGen at the balanced quality preset took between "
    "five and a half and seven minutes per request. The gTTS voiceover "
    "stage completed in approximately two seconds across all three runs. "
    "The video-assembly stage, which includes three Pollinations.ai scene "
    "calls plus FFmpeg encoding and audio overlay, took approximately "
    "three minutes regardless of the prompt content."
)
add_results_table(
    temp,
    table_number=3,
    caption="Stage-level latency across the three benchmark scenarios",
)
add_para(temp,
    "The aggregate end-to-end pipeline latency is the sum of one "
    "music-generation stage, the voiceover stage, and the video-assembly "
    "stage. A free-tier user invoking the Pulsefy AI engine reaches a "
    "finished ad in approximately 3 minutes 11 seconds on average across "
    "the three scenarios. A premium user invoking the MusicGen engine "
    "reaches a finished ad in approximately 9 minutes 12 seconds. Both "
    "totals fall within the order-of-magnitude target set in Chapter 1, "
    "where a manual workflow across separate design, music, voiceover, "
    "and editing tools was estimated at approximately eight hours. The "
    "free-tier number translates to roughly a 150-fold reduction in "
    "wall-clock production time compared to the manual baseline; the "
    "premium-tier number, slower because of MusicGen's longer inference, "
    "still represents an approximately 52-fold reduction."
)
add_para(temp,
    "Several caveats apply to these measurements. The sample size is "
    "small (three trials per stage), the runs were taken on a single "
    "hardware configuration and a single network connection, and no "
    "assessment of the qualitative output was made: the benchmark "
    "establishes that the pipeline completes within the claimed time "
    "budget, not that the produced ads are competitive with "
    "professional output. MusicGen latency was measured with a warm "
    "model cache; a cold-start invocation, in which the audiocraft "
    "library loads the checkpoint into memory, adds approximately "
    "thirty seconds to the first request after the server boots. "
    "Pollinations.ai latency depends on the load on its public open API "
    "at the moment of the request, which varies by hour."
)
add_para(temp,
    "Figure 58 presents the same data as a stacked horizontal bar, "
    "showing how each of the four stages contributes to the total "
    "pipeline latency on the free and premium tiers."
)
add_figure_placeholder(
    temp,
    58,
    "Pulsefy ad-creation pipeline latency by stage (free vs premium tier)",
)


main = Document(DOCX)
main_body = main.element.body

start_elem = None
end_elem = None
for child in list(main_body):
    if not child.tag.endswith("}p"):
        continue
    text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
    if start_elem is None and text.strip().startswith("6 Validation"):
        start_elem = child
    elif start_elem is not None and end_elem is None and text.strip().startswith("Conclusions"):
        end_elem = child
        break

assert start_elem is not None, "could not find '6 Validation' boundary"
assert end_elem is not None, "could not find 'Conclusions' boundary"

to_remove = []
collecting = False
for child in list(main_body):
    if child is start_elem:
        collecting = True
    if child is end_elem:
        break
    if collecting:
        to_remove.append(child)
for elem in to_remove:
    main_body.remove(elem)

end_idx = list(main_body).index(end_elem)
for child in list(temp.element.body):
    if child.tag.endswith("}sectPr"):
        continue
    main_body.insert(end_idx, deepcopy(child))
    end_idx += 1

main.save(DOCX)
print(f"Saved: {DOCX}")
print(f"Removed {len(to_remove)} old elements.")
