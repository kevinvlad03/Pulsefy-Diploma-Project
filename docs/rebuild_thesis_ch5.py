"""
Rebuild Chapter 5 (Implementation) of Pulsefy_Diploma_Thesis.docx.

New structure:
- 5.1 Technologies (Development, External services)
- 5.2 AI generation subsystems with code-showcase-first approach
- 5.3 Web application walkthrough (Advertiser flow + Creator leaderboard;
                                    no admin panel)

Figure-placeholder caption convention:
- "Code excerpt: <file path>" for code screenshots Vlad will paste
- "Screenshot: <view>" for app screenshots Vlad will paste
- Decorative figures (logo ribbon, env file) keep their plain name

Run from the worktree root:
    python3 docs/rebuild_thesis_ch5.py
"""

from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def add_heading_styled(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = "Times New Roman"
    return h


def add_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_figure_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[FIGURE {number}: {caption} — to be inserted]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {number}: {caption}")
    run.font.size = Pt(10)


def add_table_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[TABLE {number}: {caption} — to be filled]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Table {number}: {caption}")
    run.font.size = Pt(10)


temp = Document()

# ─── 5 Implementation ───────────────────────────────────────────────────────
add_heading_styled(temp, "5 Implementation", level=1)

# ─── 5.1 Technologies ───────────────────────────────────────────────────────
add_heading_styled(temp, "5.1 Technologies", level=2)
add_heading_styled(temp, "5.1.1 Development", level=3)
add_para(temp,
    "Pulsefy's front-end is a single-page React 18 application written in "
    "TypeScript and bundled with Vite. The component layer uses shadcn/ui "
    "primitives styled with Tailwind CSS, the data layer uses TanStack Query "
    "for cache-aware HTTP calls against the backend, and React Router v6 "
    "drives the navigation. State that crosses route boundaries (the active "
    "audio playback, the current subscription tier, the toast queue) lives "
    "in lightweight context providers rather than a global store."
)
add_para(temp,
    "The back-end is a Node.js 18 Express 4 server written as an ESM "
    "module. The HTTP layer uses Zod for runtime payload validation, "
    "bcryptjs for password hashing at cost 10, and jsonwebtoken for HS256 "
    "session tokens with a seven-day expiry. Database access is direct "
    "through the pg library; there is no ORM, and the schema is applied "
    "from idempotent SQL scripts at server start. Each route file is wrapped "
    "with an async handler utility that forwards rejected promises to "
    "Express error middleware."
)
add_para(temp,
    "The AI and media generation stack runs as separate processes. A "
    "Python 3.11 virtual environment under server/musicgen/.venv holds the "
    "pinned dependencies for the audiocraft library that powers MusicGen, "
    "the gTTS client, and the custom PyTorch model used by the Pulsefy AI "
    "free engine. FFmpeg is invoked separately as a system binary for video "
    "encoding and audio overlay. Each invocation spawns a fresh subprocess "
    "from Node via child_process.spawn and returns either a JSON result or "
    "the path of the generated media file. Figure 40 summarises the stack."
)
add_figure_placeholder(temp, 40, "Technology stack (logo ribbon)")

add_heading_styled(temp, "5.1.2 External services", level=3)
add_para(temp,
    "Three external services are reached over HTTPS. Pollinations.ai "
    "provides text-to-image inference under an MIT-licensed open API and is "
    "used for ad scene generation. No API key is required for the public "
    "tier; the service rate-limits bursts, which the imagegen.js client "
    "respects with a three-second inter-scene sleep and a three-attempt "
    "retry policy with exponential backoff."
)
add_para(temp,
    "Jamendo's v3.0 catalog API supplies the licensed-music browser used "
    "for inspiration and for direct attachment to ad projects. Calls are "
    "proxied through the server-side /jamendo route to keep the client_id "
    "out of the browser. When the upstream API is unavailable or returns an "
    "empty set, the proxy falls back to a locally seeded mini-catalog so "
    "the user-facing experience remains intact."
)
add_para(temp,
    "The MusicGen weights are fetched from the Hugging Face model hub on "
    "first run of the Python subprocess. The small variant is 600 megabytes "
    "and is cached under the user's Hugging Face home directory; subsequent "
    "invocations skip the download. The free Pulsefy AI engine is a "
    "self-contained PyTorch checkpoint shipped under server/musicgen/lstm/"
    "checkpoints, which removes the network dependency for free-tier users."
)

# ─── 5.2 AI generation subsystems ───────────────────────────────────────────
add_heading_styled(temp, "5.2 AI generation subsystems", level=2)
add_para(temp,
    "Pulsefy provides four distinct AI generation subsystems. Each one is "
    "exposed through a route on the /ai prefix, persists its output to "
    "public/media/generated/, and writes a record to its dedicated table "
    "(ai_generations, tts_generations, video_generations) for retrieval. "
    "Table 2 summarises the public functions each subsystem exposes; the "
    "subsections that follow describe the implementation, with code "
    "excerpts taken directly from the production codebase."
)
add_table_placeholder(temp, 2, "AI generation subsystems and their public functions")

add_heading_styled(temp, "5.2.1 Music generation (Pulsefy AI and MusicGen)", level=3)
add_para(temp,
    "The music subsystem is the most architecturally substantial part of "
    "Pulsefy and the only one that ships two interchangeable engines. The "
    "free tier uses Pulsefy AI, a custom LSTM-based PyTorch model trained "
    "from scratch on a curated MIDI dataset. The premium tier uses Meta's "
    "MusicGen via the audiocraft library. The choice between engines is "
    "made server-side based on the user's subscription_tier and the optional "
    "modelChoice field in the request payload; a free user who requests "
    "MusicGen receives a 402 Premium Required response."
)
add_para(temp,
    "Pulsefy AI is implemented as a token-level next-event predictor over a "
    "vocabulary of MIDI events. The vocabulary contains 164 tokens that "
    "encode note pitch, duration, and rest events. The MusicLSTM model "
    "consists of an embedding layer with 64-dimensional vectors, two stacked "
    "LSTM layers with 256 hidden units and a dropout of 0.3, and a linear "
    "projection back to the vocabulary. The total parameter count is "
    "approximately 908 thousand, which fits comfortably on Apple Metal (MPS) "
    "consumer hardware and CPUs alike. Figure 41 shows the relevant section "
    "of model.py."
)
add_figure_placeholder(temp, 41, "Code excerpt: MusicLSTM model definition (server/musicgen/lstm/model.py)")
add_para(temp,
    "Training is performed by lstm/train.py against a corpus of 1034 MIDI "
    "files (Nottingham folk waltzes and Christmas tunes), parsed by "
    "lstm/parser.py and tokenised by lstm/vocab.py and lstm/dataset.py. The "
    "preprocessing pipeline yields 573,692 tokens that are split into "
    "443,132 fixed-length training windows, with a 95-to-5 train-validation "
    "split (420,976 training windows and 22,156 validation windows). The "
    "training script iterates for 30 epochs with a cosine-annealed learning "
    "rate starting at one times ten to the minus three, saves the best "
    "model on every validation-loss improvement, and exports the final "
    "checkpoint under lstm/checkpoints. Figure 42 shows the training-loop "
    "core, and Figure 43 captures one full training run."
)
add_figure_placeholder(temp, 42, "Code excerpt: LSTM training loop (server/musicgen/lstm/train.py)")
add_figure_placeholder(temp, 43, "Screenshot: Pulsefy AI training session output (Apple MPS, 30 epochs)")
add_para(temp,
    "The premium-tier MusicGen engine is invoked through "
    "server/src/services/musicgen.js, which spawns a fresh Python subprocess "
    "for every generation. The subprocess loads the requested model variant "
    "(small, medium, or large), generates the requested duration, writes a "
    "16-bit PCM WAV file under public/media/generated/, and prints a JSON "
    "line on stdout. The Node service waits for the subprocess to exit, "
    "parses the JSON, and returns the audio URL to the caller. Both engines "
    "share the same persistence flow: a pending row is inserted into "
    "ai_generations before generation starts, the row is updated to "
    "completed with the resulting audio_url on success, or to failed with "
    "an error message in the lyrics column on failure."
)
add_figure_placeholder(temp, 44, "Code excerpt: tier-routed dispatch in server/src/routes/ai.js")

add_heading_styled(temp, "5.2.2 Voiceover generation (gTTS)", level=3)
add_para(temp,
    "The voiceover subsystem wraps the gTTS Python library, which "
    "communicates with the unofficial Google Translate text-to-speech "
    "endpoint and emits an MP3 stream. The Pulsefy implementation invokes "
    "gTTS, then runs ffmpeg in a second step to transcode the MP3 to "
    "44.1 kHz mono 16-bit PCM WAV, which is the format the rest of the "
    "platform expects for audio assets. The WAV is written under "
    "public/media/generated/ with a tts- prefix to distinguish it from "
    "MusicGen outputs."
)
add_para(temp,
    "Subscription gating is enforced before the subprocess is launched: "
    "if the request language is anything other than English and the user "
    "is on the free tier, the route returns a 402 Premium Required "
    "response without consuming the gTTS quota. A successful generation "
    "writes a row to the tts_generations table containing the original "
    "script, the language code, the resulting audio URL, and the probed "
    "duration in seconds. Figure 45 shows the relevant section of the "
    "Python runner."
)
add_figure_placeholder(temp, 45, "Code excerpt: gTTS pipeline (server/tts/generate.py)")

add_heading_styled(temp, "5.2.3 Scene image generation (Pollinations.ai)", level=3)
add_para(temp,
    "Scene images for ad videos are produced by the imagegen.js service, "
    "which builds a prompt by concatenating the product name, a product "
    "description, a scene angle drawn from a fixed list (hero shot, "
    "lifestyle context, close-up detail, brand moment), and a style "
    "modifier (minimalist, vibrant, cinematic, corporate, or lifestyle). "
    "The combined prompt is URL-encoded and submitted to Pollinations.ai "
    "with width and height derived from the requested aspect ratio: 9:16 "
    "produces 720 by 1280, 1:1 produces 1080 by 1080, and 16:9 produces "
    "1280 by 720."
)
add_para(temp,
    "The HTTP client retries up to three times. On a 429 response it waits "
    "five seconds multiplied by the attempt number; on any other thrown "
    "error it waits three seconds multiplied by the attempt number. A "
    "response shorter than 1000 bytes is treated as a Pollinations error "
    "page rather than image data. Between scenes the client sleeps three "
    "seconds to avoid bursting the upstream rate limit. Figure 46 shows the "
    "prompt assembly and retry logic in their final form."
)
add_figure_placeholder(temp, 46, "Code excerpt: Pollinations.ai client (server/src/services/imagegen.js)")

add_heading_styled(temp, "5.2.4 Video assembly (FFmpeg)", level=3)
add_para(temp,
    "Pulsefy's video pipeline has two entry points. The first is the "
    "generative path: the user supplies a product brief, the videogen.js "
    "service generates the scene images, applies a Ken Burns slow-zoom "
    "effect through FFmpeg, and concatenates the resulting clips into a "
    "single video. If the user also picked a background music track and "
    "a voiceover, the service overlays both audio sources with the "
    "voiceover at full volume and the background music ducked to "
    "approximately 35 percent. The result is encoded as an MP4 at the "
    "user-selected aspect ratio."
)
add_para(temp,
    "The second entry point is the upload path: the user supplies a "
    "finished video file and selects audio sources from their existing "
    "Pulsefy assets. The videoupload.js service validates the format and "
    "size (MP4, MOV, WebM, or AVI up to 100 MB), then invokes FFmpeg with "
    "the -map flags that drop the uploaded video's original audio track "
    "and replace it with the chosen combination. Both paths persist their "
    "result to the video_generations table with a source field that "
    "distinguishes generated from uploaded. Figure 47 shows the FFmpeg "
    "command construction in videogen.js."
)
add_figure_placeholder(temp, 47, "Code excerpt: FFmpeg orchestration (server/src/services/videogen.js)")

add_heading_styled(temp, "5.2.5 Configuration", level=3)
add_para(temp,
    "All runtime configuration is driven from a single .env file at the "
    "server root. The required keys are DATABASE_URL for the PostgreSQL "
    "connection string, JWT_SECRET for session-token signing, "
    "JAMENDO_CLIENT_ID for the catalog proxy, MUSICGEN_PYTHON for the path "
    "to the Python interpreter that hosts audiocraft, MUSICGEN_MODEL for "
    "the default MusicGen variant (typically facebook/musicgen-small), "
    "MUSICGEN_DEVICE to select CPU, CUDA, or MPS, and PORT for the Express "
    "listener. Optional keys override file-system paths for the script "
    "runners and output directories. The .env file is excluded from version "
    "control; an .env.example template is committed for onboarding. "
    "Figure 48 reproduces the template."
)
add_figure_placeholder(temp, 48, "Code excerpt: .env.example template")

# ─── 5.3 Web application walkthrough ───────────────────────────────────────
add_heading_styled(temp, "5.3 Web application walkthrough", level=2)
add_para(temp,
    "This section walks through the application from an advertiser's "
    "perspective, from initial sign-up through the final ad export. A "
    "short second subsection (5.3.2) demonstrates the creator leaderboard "
    "as the public face of the platform's social layer. Administrative "
    "tooling is intentionally not included; the only privileged user role "
    "in the current build is the self-service premium upgrade described "
    "in UC11."
)

add_heading_styled(temp, "5.3.1 Walkthrough for an advertiser", level=3)
add_para(temp,
    "The walkthrough begins on the public landing page. A new visitor "
    "clicks Sign up, fills in email, display name, and password, and is "
    "redirected to the dashboard with a valid session. Returning users land "
    "on the same login form. Figure 49 shows both flows. After "
    "authentication, every page in the application is rendered inside the "
    "Layout shell, which embeds the persistent navigation rail and the "
    "global audio player at the bottom of the viewport."
)
add_figure_placeholder(temp, 49, "Screenshot: Pulsefy – Login and signup")
add_para(temp,
    "The Sound Studio is the central authoring surface and is reached "
    "either from the home page or from the navigation rail. It exposes "
    "four tabs (Music, Voiceover, Video, Upload) that mirror the four AI "
    "subsystems documented in section 5.2. The Music tab, shown in Figure "
    "50, lets the user pick between the free Pulsefy AI engine and the "
    "premium MusicGen model. Free users see the MusicGen card greyed out "
    "with a Premium label; clicking it opens the upgrade modal rather than "
    "submitting a paid request that would 402."
)
add_figure_placeholder(temp, 50, "Screenshot: Sound Studio – Music tab with Pulsefy AI vs MusicGen selector")
add_para(temp,
    "The Voiceover tab takes a script up to two thousand characters, a "
    "target language picked from the gTTS-supported list, and an optional "
    "slow-speech toggle. Free users are restricted to English; selecting "
    "another language without an active premium subscription surfaces an "
    "upgrade prompt and blocks the request. Figure 51 shows the voiceover "
    "panel and the recent-voiceovers history that lets the user audition, "
    "rename, and delete previous outputs."
)
add_figure_placeholder(temp, 51, "Screenshot: Sound Studio – Voiceover tab")
add_para(temp,
    "The Video tab is the end-to-end ad generator. The user provides a "
    "product name, a longer description, a visual style, an aspect ratio, "
    "and a scene count between one and four (capped at two for free "
    "users), and optionally attaches a previously-generated music track "
    "and a previously-generated voiceover. The request returns immediately "
    "with a 202 status and a pending record; the panel polls the videos "
    "list every few seconds until the status flips to completed. Figure "
    "52 captures the form and the inline pending state."
)
add_figure_placeholder(temp, 52, "Screenshot: Sound Studio – Video tab with ad video form")
add_para(temp,
    "The Upload tab serves the alternative path for advertisers who have "
    "already produced their own video and want only the audio overlay. "
    "A drag-and-drop zone accepts MP4, MOV, WebM, or AVI files up to one "
    "hundred megabytes, and two dropdowns let the user pick the music and "
    "voiceover tracks from their previous Pulsefy generations. The server "
    "replaces the uploaded video's original audio track with the selected "
    "combination using FFmpeg's -map flags and surfaces the processed file "
    "in the same Uploaded Videos panel shown in Figure 53."
)
add_figure_placeholder(temp, 53, "Screenshot: Sound Studio – Upload tab and processed uploads panel")
add_para(temp,
    "Outside the Sound Studio, the Jamendo music inspiration browser "
    "offers a Creative-Commons-licensed catalog that creators can preview "
    "and attach as a music asset without paying licensing fees. The browser "
    "supports search by track name or artist and filtering by genre, and "
    "every track exposes its licence URL so the chosen attribution can be "
    "carried into the final ad. Figure 54 shows the browser with the "
    "filter sidebar open."
)
add_figure_placeholder(temp, 54, "Screenshot: Jamendo music inspiration browser")
add_para(temp,
    "Once an ad is generated, it lands in the user's gallery alongside "
    "every previous generation. Each card carries a status badge "
    "(processing, completed, failed), inline playback controls, and an "
    "action menu that lets the user rename, delete, or download the file. "
    "Figure 55 shows the gallery. The download action serves the encoded "
    "MP4 directly from public/media/generated/ via the platform's static "
    "media route; no extra server processing is required at export time."
)
add_figure_placeholder(temp, 55, "Screenshot: Generated ad gallery with status badges and download actions")
add_para(temp,
    "Figure 56 closes the walkthrough by showing the final exported ad "
    "ready for upload to TikTok, Instagram Reels, or YouTube Shorts. The "
    "encoded file is delivered at the aspect ratio originally requested in "
    "the Video tab, with the correct dimensions for the target platform "
    "and the chosen audio overlaid in place of any original soundtrack."
)
add_figure_placeholder(temp, 56, "Screenshot: Exported ad ready for download")

add_heading_styled(temp, "5.3.2 Creator leaderboard", level=3)
add_para(temp,
    "The creator leaderboard is reached from the Social page. It ranks the "
    "top twenty users on the platform by total AI generation count, with "
    "follower count as a tie-breaker. The aim is to give active advertisers "
    "a public destination for their work and to make the platform's social "
    "layer visible without forcing every user into a follow graph. Each "
    "row shows the user's display name, bio, generation count, and follower "
    "count, and is clickable to open that user's public profile."
)
add_para(temp,
    "Internally, the leaderboard is a single SQL query that left-joins "
    "users to a follower-count subquery and an aggregated generation-count "
    "expression, orders the result by generation count descending and "
    "follower count descending, and limits to twenty rows. The query runs "
    "on demand when the user opens the tab; results are not pre-computed "
    "or cached. Figure 57 shows the leaderboard view as it appears to a "
    "logged-in user."
)
add_figure_placeholder(temp, 57, "Screenshot: Creator leaderboard view in Social page")


# ─── splice into main ───────────────────────────────────────────────────────
main = Document(DOCX)
main_body = main.element.body

start_elem = None
end_elem = None
for child in list(main_body):
    if not child.tag.endswith("}p"):
        continue
    text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
    if start_elem is None and text.strip().startswith("5 Implementation"):
        start_elem = child
    elif start_elem is not None and end_elem is None and text.strip().startswith("6 Validation"):
        end_elem = child
        break

assert start_elem is not None, "could not find '5 Implementation' boundary"
assert end_elem is not None, "could not find '6 Validation' boundary"

# Remove old Ch 5 elements
to_remove = []
collecting = False
for child in list(main_body):
    if child is start_elem:
        collecting = True
    if child is end_elem:
        break
    if collecting:
        to_remove.append(child)
for elem in to_remove:
    main_body.remove(elem)

# Insert temp Ch 5 content before '6 Validation'
end_idx = list(main_body).index(end_elem)
for child in list(temp.element.body):
    if child.tag.endswith("}sectPr"):
        continue
    main_body.insert(end_idx, deepcopy(child))
    end_idx += 1

main.save(DOCX)
print(f"Saved: {DOCX}")
print(f"Removed {len(to_remove)} old elements.")
