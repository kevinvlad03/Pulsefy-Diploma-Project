"""
Fill the remaining back matter of Pulsefy_Diploma_Thesis.docx:
  - Conclusions (5-block IARCA structure)
  - References (34 numbered entries)
  - Glossary additions (4 new terms relevant to the ad-creation framing)

Run from the worktree root:
    python3 docs/rebuild_thesis_back_matter.py
"""

from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def add_para(doc, text, size=11, italic=False, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    p.alignment = alignment
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_ref(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"[{n}] {text}")
    run.font.size = Pt(10)
    return p


# ─── Conclusions block ─────────────────────────────────────────────────────
conc_temp = Document()
add_para(conc_temp,
    "This thesis presented Pulsefy, an AI-assisted ad creation platform "
    "that bundles four AI generation subsystems (text-to-image scene "
    "generation, text-to-music generation with a free LSTM engine and a "
    "premium MusicGen engine, text-to-speech voiceover, and FFmpeg-based "
    "video assembly) behind a single web application aimed at solo "
    "creators and small businesses producing short-form social-media "
    "advertising. The codebase, the design, and the benchmark in Chapter 6 "
    "together show that the platform produces a finished Instagram Reel "
    "ad in approximately three minutes on the free tier and nine minutes "
    "on the premium tier, compared to an estimated eight hours for the "
    "equivalent manual workflow across separate design, music, voiceover, "
    "and editing tools."
)
add_para(conc_temp, "The design was guided by the following constraints:")
add_bullet(conc_temp,
    "Accessibility for solo creators with no upfront tooling cost: the "
    "free tier provides functional coverage across all four subsystems."
)
add_bullet(conc_temp,
    "Bundling rather than fragmentation: a single login replaces four "
    "separate tool subscriptions and four interfaces."
)
add_bullet(conc_temp,
    "Honest tier separation: the premium tier unlocks higher-quality "
    "music output and multilingual voiceover, not new feature categories."
)
add_bullet(conc_temp,
    "An open-source codebase with no proprietary dependencies on the AI "
    "side beyond Meta's MusicGen, itself released under an open licence."
)
add_bullet(conc_temp,
    "Output formats matched to the dominant social-media aspect ratios "
    "(9:16, 1:1, 16:9)."
)
add_bullet(conc_temp,
    "Synchronous AI generation with clear progress feedback rather than "
    "opaque background queues."
)
add_para(conc_temp,
    "I implemented Pulsefy with these constraints in mind over the course "
    "of the academic year, training the Pulsefy AI LSTM engine on a corpus "
    "of 1034 MIDI files to give free-tier users an instant alternative to "
    "MusicGen, and stitching the four AI subsystems into a uniform request-"
    "response interface so that adding a new subsystem in the future would "
    "not require re-architecting the rest."
)
add_para(conc_temp,
    "Several limitations remain. The benchmark in Chapter 6 used a sample "
    "size of three trials per stage on a single hardware configuration and "
    "did not include a qualitative assessment of the generated content; a "
    "larger-scale evaluation with a listening study (N ≥ 20) and an "
    "objective audio-distance metric such as Fréchet Audio Distance would "
    "strengthen the quality claim. The MusicGen subprocess remains the "
    "slowest stage of the pipeline (between five and seven minutes on "
    "consumer CPU), which a job-queue-backed worker pool with GPU "
    "acceleration would resolve at production scale. Subscription payments "
    "are simulated rather than integrated with a real payment processor. "
    "The four-tab Sound Studio currently expects a desktop-class viewport; "
    "a mobile-responsive layout would broaden the addressable user base."
)
add_para(conc_temp,
    "The intersection of generative AI and short-form social-media "
    "advertising is moving fast. Pulsefy aims to be a small, honest "
    "contribution to the open-source side of it."
)


# ─── References block ──────────────────────────────────────────────────────
refs_temp = Document()
refs = [
    "Adobe. (2026). *Creative Cloud Pro*. Adobe. Retrieved May 7, 2026, from https://www.adobe.com/creativecloud/pro.html",
    "Agostinelli, A., Denk, T. I., Borsos, Z., Engel, J., Verzetti, M., Caillon, A., ... & Frank, C. (2023). MusicLM: Generating Music From Text. arXiv:2301.11325. https://arxiv.org/abs/2301.11325",
    "audiocraft. (2024). *Audiocraft: A library for audio processing and generation with deep learning* [Source code]. Meta AI. GitHub. Retrieved May 7, 2026, from https://github.com/facebookresearch/audiocraft",
    "Black Forest Labs. (2024, August 1). *Announcing Black Forest Labs*. Black Forest Labs. https://bfl.ai/announcing-black-forest-labs/",
    "Canva. (2025). *Canva 2025 Wrap*. Canva Newsroom. Retrieved May 5, 2026, from https://www.canva.com/newsroom/news/canva-2025-wrap/",
    "Copet, J., Kreuk, F., Gat, I., Remez, T., Kant, D., Synnaeve, G., Adi, Y., & Défossez, A. (2024). Simple and Controllable Music Generation. arXiv:2306.05284. https://arxiv.org/abs/2306.05284",
    "Express. (2024). *Express — Node.js web application framework*. OpenJS Foundation. Retrieved May 7, 2026, from https://expressjs.com/",
    "FFmpeg. (2024). *FFmpeg Documentation*. Retrieved May 7, 2026, from https://ffmpeg.org/documentation.html",
    "gTTS. (2024). *gTTS: Google Text-to-Speech* [Software]. PyPI. https://pypi.org/project/gTTS/",
    "Ho, J., Jain, A., & Abbeel, P. (2020, June 19). Denoising Diffusion Probabilistic Models. arXiv:2006.11239. https://arxiv.org/abs/2006.11239",
    "IFPI. (2026, March). *Global Music Report 2026*. International Federation of the Phonographic Industry. https://www.ifpi.org/global-music-report-2026-global-recorded-music-revenues-grow-6-4-as-record-companies-drive-innovation/",
    "Jamendo. (n.d.). *Jamendo Developer — API v3.0 documentation*. Retrieved May 5, 2026, from https://developer.jamendo.com/v3.0",
    "Jones, M., Bradley, J., & Sakimura, N. (2015, May). JSON Web Token (JWT). IETF RFC 7519. https://www.rfc-editor.org/rfc/rfc7519",
    "Liu, H., Chen, Z., Yuan, Y., Mei, X., Liu, X., Mandic, D., Wang, W., & Plumbley, M. D. (2023). AudioLDM: Text-to-Audio Generation with Latent Diffusion Models. arXiv:2301.12503. https://arxiv.org/abs/2301.12503",
    "Music Business Worldwide. (2024). *There are now 120,000 new tracks hitting music streaming services each day*. https://www.musicbusinessworldwide.com/there-are-now-120000-new-tracks-hitting-music-streaming-services-each-day/",
    "Node.js. (2024). *Node.js Documentation*. OpenJS Foundation. Retrieved May 7, 2026, from https://nodejs.org/en/docs/",
    "OpenAI. (2024). *InVideo AI — Customer Story*. OpenAI. Retrieved May 5, 2026, from https://openai.com/index/invideo-ai/",
    "Pollinations. (2026). *Pollinations* [Source code]. GitHub. Retrieved May 9, 2026, from https://github.com/pollinations/pollinations",
    "PostgreSQL Global Development Group. (2024). *PostgreSQL 14 Documentation*. https://www.postgresql.org/docs/14/",
    "PyTorch. (2024). *PyTorch Documentation*. The Linux Foundation. https://pytorch.org/docs/stable/index.html",
    "React. (2024). *React Documentation*. Meta. https://react.dev/",
    "Rombach, R., Blattmann, A., Lorenz, D., Esser, P., & Ommer, B. (2022). High-Resolution Image Synthesis with Latent Diffusion Models. *Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR 2022)*. https://arxiv.org/abs/2112.10752",
    "shadcn. (2024). *shadcn/ui*. Retrieved May 7, 2026, from https://ui.shadcn.com/",
    "Sortlist. (2026). *The 10 best advertising agencies in Romania*. Sortlist. Retrieved May 7, 2026, from https://www.sortlist.com/advertising/romania-ro",
    "Splice. (2024). *Which video editing apps have the largest user base?* Splice Blog. Retrieved May 5, 2026, from https://spliceapp.com/blog/which-video-editing-apps-have-the-largest-user-base/",
    "Spotify. (2026, April 28). *Spotify reports first quarter 2026 earnings*. Spotify Newsroom. https://newsroom.spotify.com/2026-04-28/spotify-q1-2026-earnings/",
    "Spotify. (2025, June 30). *Discover Weekly turns 10*. Spotify Newsroom. https://newsroom.spotify.com/2025-06-30/discover-weekly-turns-10-celebrating-100-billion-tracks-streamed-and-a-decade-of-personalized-discovery/",
    "Stability AI. (2024, April 3). *Stable Audio 2.0*. Stability AI. https://stability.ai/news/stable-audio-2-0",
    "Statista. (2025). *Number of monthly active TikTok users worldwide from 2018 to 2025*. Statista. Retrieved May 7, 2026, from https://www.statista.com/statistics/1327116/number-of-global-tiktok-users/",
    "Tailwind Labs. (2024). *Tailwind CSS Documentation*. https://tailwindcss.com/docs",
    "TanStack. (2024). *TanStack Query Documentation*. https://tanstack.com/query/latest",
    "Vite. (2024). *Vite Documentation*. https://vitejs.dev/",
    "Williams, R. (2025, September 25). *Global ad spending now estimated to rise 7.4%, WARC's first positive revision in over a year*. EMARKETER. https://www.emarketer.com/content/global-ad-spending-now-estimated-rise-7-4--warc-s-first-positive-revision-over-year",
    "Zod. (2024). *Zod: TypeScript-first schema validation*. https://zod.dev/",
]
for i, ref in enumerate(refs, start=1):
    add_ref(refs_temp, i, ref)


# ─── Splice into main ───────────────────────────────────────────────────────
main = Document(DOCX)
main_body = main.element.body


def replace_section_content(start_heading_text, end_heading_text_or_none, temp_doc):
    """Replace all body content from the heading immediately AFTER
    start_heading_text up to (but not including) end_heading_text.
    If end_heading_text_or_none is None, deletes through the end of the body
    (excluding the trailing sectPr).
    """
    start_elem = None
    end_elem = None
    for child in list(main_body):
        if not child.tag.endswith("}p"):
            continue
        text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
        if start_elem is None and text.strip().startswith(start_heading_text):
            start_elem = child
            continue
        if start_elem is not None and end_heading_text_or_none and text.strip().startswith(end_heading_text_or_none):
            end_elem = child
            break

    assert start_elem is not None, f"could not find '{start_heading_text}' heading"

    # Remove everything between start_elem (exclusive) and end_elem (exclusive)
    to_remove = []
    collecting = False
    for child in list(main_body):
        if child is start_elem:
            collecting = True
            continue
        if end_elem is not None and child is end_elem:
            break
        if collecting:
            # stop if we hit a sectPr (end of body)
            if child.tag.endswith("}sectPr"):
                break
            to_remove.append(child)

    for elem in to_remove:
        main_body.remove(elem)

    # Insert temp content right after start_elem
    insert_idx = list(main_body).index(start_elem) + 1
    for child in list(temp_doc.element.body):
        if child.tag.endswith("}sectPr"):
            continue
        main_body.insert(insert_idx, deepcopy(child))
        insert_idx += 1


# Conclusions
replace_section_content("Conclusions", "References", conc_temp)
print("Conclusions: filled")

# References
replace_section_content("References", "Glossary", refs_temp)
print(f"References: {len(refs)} entries inserted")


# Glossary additions: find the glossary table, append 4 new terms
glossary_additions = [
    ("DDPM", "Denoising Diffusion Probabilistic Model"),
    ("LSTM", "Long Short-Term Memory (a recurrent neural network architecture)"),
    ("MIDI", "Musical Instrument Digital Interface"),
    ("MPS", "Metal Performance Shaders (Apple's GPU compute API)"),
]
glossary_table = None
for t in main.tables:
    cells = [r.cells[0].text for r in t.rows]
    if "API" in cells and "JSON" in cells:
        glossary_table = t
        break
if glossary_table is not None:
    for term, definition in glossary_additions:
        new_row = glossary_table.add_row()
        new_row.cells[0].text = term
        for p in new_row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        new_row.cells[1].text = definition
        for p in new_row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
        new_row.cells[0].width = Cm(3.5)
        new_row.cells[1].width = Cm(13)
    print(f"Glossary: {len(glossary_additions)} entries appended")
else:
    print("WARN: glossary table not found, skipping additions")


main.save(DOCX)
print(f"Saved: {DOCX}")
