"""
Targeted edit of Chapter 1 in Pulsefy_Diploma_Thesis.docx.
Replaces unsourced / outdated claims with verified citations from the research bank.

Does NOT regenerate the whole document — only modifies three specific paragraphs
and removes their two corresponding [TODO] markers.

Run from the worktree root:
    python3 docs/edit_thesis_ch1_citations.py
"""

from docx import Document
from docx.shared import Pt, RGBColor

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def find_paragraph_index(doc, signature):
    for i, p in enumerate(doc.paragraphs):
        if signature in p.text:
            return i
    return -1


def replace_first_run(p, new_text, fontsize=11):
    """Set the first run's text and strip any subsequent runs."""
    if p.runs:
        p.runs[0].text = new_text
        for run in list(p.runs[1:]):
            run._element.getparent().remove(run._element)
    else:
        run = p.add_run(new_text)
        run.font.size = Pt(fontsize)


def append_marker(p, marker_text):
    """Append a red-italic [marker] run, signalling 'add Word footnote here'."""
    run = p.add_run(f"  [{marker_text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.font.size = Pt(10)


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


doc = Document(DOCX)

# ----- 1.1 paragraph 1: Spotify MAU + IFPI macro -----
i = find_paragraph_index(doc, "Spotify reporting over 600 million")
assert i >= 0, "could not find 1.1 P1 (Spotify MAU paragraph)"
p = doc.paragraphs[i]
replace_first_run(
    p,
    "Music is one of the most consumed forms of media. Spotify alone reported "
    "761 million monthly active users and 293 million Premium subscribers in the "
    "first quarter of 2026, and the global recorded music market reached "
    "US$31.7 billion in 2025, with streaming accounting for 69.6 percent of "
    "revenue.",
)
append_marker(
    p,
    "footnotes: Spotify Newsroom Q1 2026; IFPI Global Music Report 2026 — "
    "insert Word footnotes here",
)

# remove the [TODO: cite the specific Spotify quarterly investor report] paragraph that follows
todo_p = doc.paragraphs[i + 1]
if "[TODO:" in todo_p.text and "Spotify quarterly" in todo_p.text:
    remove_paragraph(todo_p)
else:
    print(f"WARN: expected TODO paragraph at index {i+1}, got: {todo_p.text[:80]!r}")

# ----- 1.1 paragraph 2: catalog growth + uploads/day -----
i = find_paragraph_index(doc, "The catalog grows constantly")
assert i >= 0, "could not find 1.1 P2 (catalog growth paragraph)"
p = doc.paragraphs[i]
replace_first_run(
    p,
    "The catalog grows constantly. Around 120,000 new tracks reach streaming "
    "services every day, and any user with a laptop and free recording software "
    "can publish music that previously required a record label or studio access.",
)
append_marker(
    p,
    "footnote: Music Business Worldwide, citing Luminate 2025 — "
    "insert Word footnote here",
)

# ----- 1.2 paragraph 1: recommendation engagement (Discover Weekly) -----
i = find_paragraph_index(doc, "more relevant content and more variety")
assert i >= 0, "could not find 1.2 P1 (listener motivation paragraph)"
p = doc.paragraphs[i]
replace_first_run(
    p,
    "Listeners want both more relevant content and more variety from the "
    "platforms they use. Recommendation features have become central to the "
    "streaming experience: Spotify reported that its Discover Weekly playlist "
    "alone has surfaced over 100 billion tracks since launch in 2015, and "
    "users who engage with the feature stream music nearly twice as long as "
    "those who do not.",
)
append_marker(
    p,
    "footnotes: Spotify Newsroom Discover Weekly 10-year milestone, 2025; "
    "Spotify Newsroom 2020 longitudinal stat — insert Word footnotes here",
)

# remove the [TODO: cite a Spotify or Apple Music engagement report] paragraph
todo_p = doc.paragraphs[i + 1]
if "[TODO:" in todo_p.text and ("Spotify" in todo_p.text or "Apple Music" in todo_p.text):
    remove_paragraph(todo_p)
else:
    print(f"WARN: expected TODO paragraph at index {i+1}, got: {todo_p.text[:80]!r}")

doc.save(DOCX)
print(f"Saved: {DOCX}")
