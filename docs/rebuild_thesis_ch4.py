"""
Rebuild Chapter 4 (Design) of Pulsefy_Diploma_Thesis.docx with the reduced
DCD/DSD set: 6 use cases (UC2, UC5, UC6, UC7, UC8, UC11) instead of all 11.

Same splice strategy as the Ch 3 rebuild: build a temp Document with new
content, then surgically replace everything between the existing '4 Design'
heading and the existing '5 Implementation' heading.

Run from the worktree root:
    python3 docs/rebuild_thesis_ch4.py
"""

from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = (
    "/Users/dumitruvlad/Documents/Facultate/Licenta/Pulsefy/"
    ".claude/worktrees/friendly-shockley-28309b/docs/"
    "Pulsefy_Diploma_Thesis.docx"
)


def add_heading_styled(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = "Times New Roman"
    return h


def add_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_figure_placeholder(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[FIGURE {number}: {caption} — to be drawn]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {number}: {caption}")
    run.font.size = Pt(10)


temp = Document()

# ─── 4 Design ───────────────────────────────────────────────────────────────
add_heading_styled(temp, "4 Design", level=1)

# ─── 4.1 Architecture ───────────────────────────────────────────────────────
add_heading_styled(temp, "4.1 Architecture", level=2)

add_heading_styled(temp, "4.1.1 Architecture choice", level=3)
add_para(temp,
    "Pulsefy is built on a classical client-server architecture extended "
    "with separate Python and FFmpeg subprocesses for the AI generation "
    "pipeline. The web tier serves a React single-page application from a "
    "Vite development server during development and from a static build "
    "during deployment. The application tier is a Node.js Express server "
    "that exposes a JSON API, validates input with Zod schemas, and signs "
    "JWT session tokens for authenticated requests. The data tier is a "
    "single PostgreSQL database accessed through the pg library without an "
    "object-relational mapper."
)
add_para(temp,
    "A pure client-server architecture would be insufficient on its own "
    "because every AI generation in Pulsefy is computationally heavy. "
    "MusicGen inference on a CPU takes between sixty seconds and five "
    "minutes depending on the model variant; the LSTM-based Pulsefy AI "
    "engine takes several seconds; gTTS round-trips to Google's "
    "text-to-speech endpoint; and FFmpeg video assembly is bound by "
    "encoding throughput. Performing any of these inside the Node.js event "
    "loop would block all other requests for the duration of the "
    "generation."
)
add_para(temp,
    "The architecture chosen for Pulsefy keeps the Express server thin and "
    "spawns a dedicated Python or FFmpeg subprocess for every generation "
    "request. Each subprocess inherits the Node process environment, "
    "writes its output to the static media folder, and returns either a "
    "file path or a JSON payload that the Express handler relays back to "
    "the client. This keeps the Express tier responsive for non-generation "
    "traffic, isolates the heavy ML dependencies into a single Python "
    "virtual environment shared by MusicGen, the LSTM engine, and gTTS, "
    "and lets the platform expose four distinct AI subsystems (image, "
    "music, voiceover, and video assembly) through a single uniform API. "
    "Figure 22 illustrates the full data flow."
)
add_figure_placeholder(temp, 22, "Pulsefy architecture block diagram")

add_heading_styled(temp, "4.1.2 Package diagrams", level=3)
add_figure_placeholder(temp, 23, "pulsefy.server package diagram")
add_figure_placeholder(temp, 24, "pulsefy.client package diagram")

add_heading_styled(temp, "4.1.3 Deployment", level=3)
add_para(temp,
    "Pulsefy is deployed as a local development environment for the "
    "thesis demonstration. The PostgreSQL database runs locally on its "
    "default port. The Node.js application server runs on port 4000 and "
    "exposes the Express API under /auth, /ai, /jamendo, /social, "
    "/playlists, /recommendations, /listening-events, /track-likes, and "
    "/tracks. The Vite development server runs on port 8080 and serves "
    "the React SPA with hot module reload. The Python virtual environment "
    "under server/musicgen/.venv contains pinned dependencies for "
    "MusicGen, gTTS, and the LSTM training stack, and is invoked by the "
    "Node services on demand via child_process.spawn."
)
add_para(temp,
    "Production deployment is out of scope for the thesis. The "
    "architecture supports horizontal scaling of the Express tier behind "
    "a load balancer, but the synchronous AI subprocess model would "
    "require a job queue with a worker pool to handle realistic "
    "production traffic. Figure 25 shows the configuration used during "
    "the thesis demonstration."
)
add_figure_placeholder(temp, 25, "Deployment configuration")

# ─── 4.2 Detailed design ────────────────────────────────────────────────────
add_heading_styled(temp, "4.2 Detailed design", level=2)
add_para(temp,
    "The detailed design is presented in two parts. Section 4.2.1 walks "
    "through six representative use cases at the class and sequence "
    "level: the authentication baseline (UC2 Log in), the four AI "
    "generation subsystems (UC5 Generate music, UC6 Generate voiceover, "
    "UC7 Generate ad video, UC8 Upload product video and overlay audio), "
    "and the tier-management operation (UC11 Upgrade subscription tier). "
    "The five remaining use cases follow the same architectural patterns "
    "as one of these six and are not repeated. Section 4.2.2 presents "
    "the database structure from two points of view."
)

add_heading_styled(temp, "4.2.1 Design class diagrams and design sequence diagrams", level=3)

dcd_dsd_pairs = [
    ("4.2.1.1 UC2 Log in", "UC2 Log in", 26, 27),
    ("4.2.1.2 UC5 Generate music", "UC5 Generate music", 28, 29),
    ("4.2.1.3 UC6 Generate voiceover", "UC6 Generate voiceover", 30, 31),
    ("4.2.1.4 UC7 Generate ad video", "UC7 Generate ad video", 32, 33),
    ("4.2.1.5 UC8 Upload product video and overlay audio", "UC8 Upload product video and overlay audio", 34, 35),
    ("4.2.1.6 UC11 Upgrade subscription tier", "UC11 Upgrade subscription tier", 36, 37),
]
for heading, uc_name, dcd_num, dsd_num in dcd_dsd_pairs:
    add_heading_styled(temp, heading, level=3)
    add_figure_placeholder(temp, dcd_num, f"{uc_name} design class diagram")
    add_figure_placeholder(temp, dsd_num, f"{uc_name} design sequence diagram")

add_heading_styled(temp, "4.2.2 Database structure", level=3)
add_para(temp,
    "The structure of the Pulsefy database can be analysed from two "
    "complementary points of view. The first view, shown in Figure 38, "
    "isolates the AI subsystem tables (ai_generations, tts_generations, "
    "video_generations, track_likes) and the foreign keys that bind every "
    "generated asset back to its owning user. The second view, shown in "
    "Figure 39, presents the complete schema at deployment, including the "
    "authentication, social, catalog, and recommendation tables."
)
add_para(temp,
    "The schema is applied through idempotent SQL scripts executed at "
    "server start. CREATE TABLE IF NOT EXISTS and ALTER TABLE ADD COLUMN "
    "IF NOT EXISTS statements let the schema evolve without a separate "
    "migration framework. The users table carries the subscription_tier "
    "column that gates premium AI features. The ai_generations and "
    "video_generations tables persist the request payload, the chosen "
    "model, and the resulting media URL. Two many-to-many relations are "
    "realised through junction tables: playlist_tracks linking playlists "
    "and tracks, and follows linking users to other users. The "
    "playlist_likes, playlist_comments, and notifications tables support "
    "the social layer; the track_likes table records track-level "
    "appreciation."
)
add_figure_placeholder(temp, 38, "Pulsefy AI subsystem ERD")
add_figure_placeholder(temp, 39, "Pulsefy full deployment ERD")


# ─── splice into main ───────────────────────────────────────────────────────
main = Document(DOCX)
main_body = main.element.body

start_elem = None
end_elem = None
for child in list(main_body):
    if not child.tag.endswith("}p"):
        continue
    text = "".join(t.text or "" for t in child.iter() if t.tag.endswith("}t"))
    if start_elem is None and text.strip().startswith("4 Design"):
        start_elem = child
    elif start_elem is not None and end_elem is None and text.strip().startswith("5 Implementation"):
        end_elem = child
        break

assert start_elem is not None, "could not find '4 Design' boundary"
assert end_elem is not None, "could not find '5 Implementation' boundary"

# Remove old Ch 4 elements
to_remove = []
collecting = False
for child in list(main_body):
    if child is start_elem:
        collecting = True
    if child is end_elem:
        break
    if collecting:
        to_remove.append(child)
for elem in to_remove:
    main_body.remove(elem)

# Insert temp Ch 4 content before the '5 Implementation' heading
end_idx = list(main_body).index(end_elem)
for child in list(temp.element.body):
    if child.tag.endswith("}sectPr"):
        continue
    main_body.insert(end_idx, deepcopy(child))
    end_idx += 1

main.save(DOCX)
print(f"Saved: {DOCX}")
print(f"Removed {len(to_remove)} old elements.")
